from __future__ import annotations

import json
import math
import os
import subprocess
import sys
import textwrap
from datetime import datetime, timezone
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "dissertation_2_9_3"
FIG_DIR = OUT_DIR / "figures"
OUTPUT = OUT_DIR / "Dissertation_2_9_3_ACIM_AI_working_draft.docx"
EVIDENCE_DIR = OUT_DIR / "evidence"
MCSDK_PROJECT = ROOT / "mcsdk_reference" / "AIR56B2_025KW_220V_DELTA_NAMEPLATE_VF_NOT_FOR_HV"
MCSDK_BUILD_DIR = MCSDK_PROJECT / "STM32CubeIDE" / "Debug"
MCSDK_MANIFEST = MCSDK_BUILD_DIR / "ACIM-NUCLEOG431RB-IPM15B-VF_OL.build-manifest.json"
UNO_BUILD_DIR = ROOT / "firmware" / "unoq_mcsdk_scalar"
UNO_MANIFEST = UNO_BUILD_DIR / "unoq_mcsdk_scalar.build-manifest.json"
MOTOR_PROFILE = ROOT / "docs" / "mcsdk_acim_motor_profile.iek_air56b2_catalog_operator_confirmed_vf_candidate.json"
PAGE_WIDTH_DXA = 9921  # 175 mm text width: A4 minus 25/10 mm margins.

BLACK = "000000"
BLUE = "174A7E"
TEAL = "207A70"
RED = "A63A3A"
AMBER = "A06A18"
GRAY = "5B6570"
LIGHT_BLUE = "EAF1F7"
LIGHT_TEAL = "E9F3F1"
LIGHT_RED = "F8ECEC"
LIGHT_GRAY = "F2F3F4"
WHITE = "FFFFFF"

FIG_INK = "#171717"
FIG_LINE = "#4B5560"
FIG_BLUE = "#2F5D7C"
FIG_TEAL = "#3F6B62"
FIG_AMBER = "#8A6428"
FIG_RED = "#8E3B3B"
FIG_GRAY_FILL = "#F4F5F6"
FIG_BLUE_FILL = "#EEF3F6"
FIG_TEAL_FILL = "#EFF4F2"
FIG_AMBER_FILL = "#F7F3EA"
FIG_RED_FILL = "#F7EEEE"


def read_json(path: Path) -> dict:
    try:
        value = json.loads(path.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def format_bytes(value: object) -> str:
    try:
        return f"{int(value):,}".replace(",", " ")
    except (TypeError, ValueError):
        return "не подтверждено"


def artifact_sizes(manifest: dict) -> dict[str, str]:
    result: dict[str, str] = {}
    entries = manifest.get("artifacts", [])
    if not isinstance(entries, list):
        return result
    for entry in entries:
        if isinstance(entry, dict) and isinstance(entry.get("file"), str):
            result[entry["file"]] = format_bytes(entry.get("bytes"))
    return result


def run_json_check(name: str, args: list[str]) -> dict:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    completed = subprocess.run(
        [sys.executable, *args],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        check=False,
    )
    try:
        report = json.loads(completed.stdout)
    except json.JSONDecodeError:
        report = {
            "tool": name,
            "pass": False,
            "parse_error": "tool did not emit a JSON report",
            "stdout": completed.stdout,
        }
    if not isinstance(report, dict):
        report = {"tool": name, "pass": False, "parse_error": "JSON report is not an object"}
    report["runner_exit_code"] = completed.returncode
    if completed.stderr.strip():
        report["runner_stderr"] = completed.stderr.strip()
    (EVIDENCE_DIR / f"{name}.json").write_text(
        json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return report


def collect_programmatic_evidence() -> dict[str, dict]:
    reports = {
        "profile": run_json_check(
            "air56b2_firmware_profile_check",
            [str(ROOT / "tools" / "air56b2_firmware_profile_check.py")],
        ),
        "contract": run_json_check(
            "uno_nucleo_mcsdk_contract_check",
            [str(ROOT / "tools" / "uno_nucleo_mcsdk_contract_check.py")],
        ),
        "bundle": run_json_check(
            "verify_firmware_bundle",
            [str(ROOT / "tools" / "verify_firmware_bundle.py")],
        ),
    }
    release_path = EVIDENCE_DIR / "mcsdk_release_preflight.json"
    completed = subprocess.run(
        [
            sys.executable,
            str(ROOT / "tools" / "mcsdk_release_preflight.py"),
            "--project",
            str(MCSDK_PROJECT),
            "--motor-profile",
            str(MOTOR_PROFILE),
            "--artifacts",
            str(MCSDK_BUILD_DIR),
            "--output",
            str(release_path),
        ],
        cwd=ROOT,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        check=False,
    )
    release = read_json(release_path)
    if not release:
        release = {
            "tool": "mcsdk_release_preflight",
            "pass": False,
            "parse_error": "tool did not create a valid JSON report",
            "stdout": completed.stdout,
        }
    release["runner_exit_code"] = completed.returncode
    if completed.stderr.strip():
        release["runner_stderr"] = completed.stderr.strip()
    release_path.write_text(
        json.dumps(release, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    reports["release"] = release
    return reports


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths: list[int]) -> None:
    source_width = sum(widths)
    if source_width != PAGE_WIDTH_DXA:
        widths = [round(width * PAGE_WIDTH_DXA / source_width) for width in widths]
        widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    old_grid = tbl.tblGrid
    new_grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        new_grid.append(grid_col)
    tbl.replace(old_grid, new_grid)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def force_font_family(r_fonts, name: str) -> None:
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), name)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        key = qn(f"w:{attr}")
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]


def set_language(r_pr, language: str = "ru-RU") -> None:
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), language)
    lang.set(qn("w:eastAsia"), language)
    lang.set(qn("w:bidi"), language)


def force_black_color(r_pr) -> None:
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    color.set(qn("w:val"), BLACK)
    for attr in ("themeColor", "themeTint", "themeShade"):
        key = qn(f"w:{attr}")
        if key in color.attrib:
            del color.attrib[key]


def set_run_font(run, name="Times New Roman", size=14, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    force_font_family(r_pr.get_or_add_rFonts(), name)
    set_language(r_pr)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color == BLACK:
        force_black_color(r_pr)
    else:
        run.font.color.rgb = RGBColor.from_string(color)


def normalize_text_appearance(doc: Document) -> None:
    """Make font family and black text explicit instead of relying on Word themes."""
    styles_root = doc.styles._element
    defaults = styles_root.find(qn("w:docDefaults"))
    if defaults is None:
        defaults = OxmlElement("w:docDefaults")
        styles_root.insert(0, defaults)
    r_pr_default = defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        defaults.insert(0, r_pr_default)
    default_r_pr = r_pr_default.find(qn("w:rPr"))
    if default_r_pr is None:
        default_r_pr = OxmlElement("w:rPr")
        r_pr_default.append(default_r_pr)
    force_font_family(default_r_pr.get_or_add_rFonts(), "Times New Roman")
    set_language(default_r_pr)
    force_black_color(default_r_pr)

    for style in doc.styles:
        r_pr = style._element.get_or_add_rPr()
        force_font_family(r_pr.get_or_add_rFonts(), "Times New Roman")
        set_language(r_pr)
        force_black_color(r_pr)

    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
        paragraphs.extend(section.first_page_header.paragraphs)
        paragraphs.extend(section.first_page_footer.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            r_pr = run._element.get_or_add_rPr()
            force_font_family(r_pr.get_or_add_rFonts(), "Times New Roman")
            set_language(r_pr)
            force_black_color(r_pr)


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text_node, end])
    set_run_font(run, size=12)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(25)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    force_font_family(normal._element.rPr.rFonts, "Times New Roman")
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    force_font_family(h1._element.rPr.rFonts, "Times New Roman")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(18)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    force_font_family(h2._element.rPr.rFonts, "Times New Roman")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.first_line_indent = Cm(0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    force_font_family(h3._element.rPr.rFonts, "Times New Roman")
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.first_line_indent = Cm(0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.keep_with_next = True

    if "Caption" in doc.styles:
        caption = doc.styles["Caption"]
    else:
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Times New Roman"
    force_font_family(caption._element.rPr.rFonts, "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = True

    compact = doc.styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    compact.font.name = "Times New Roman"
    force_font_family(compact._element.rPr.rFonts, "Times New Roman")
    compact.font.size = Pt(12)
    compact.paragraph_format.first_line_indent = Cm(0)
    compact.paragraph_format.line_spacing = 1.0
    compact.paragraph_format.space_after = Pt(0)

    bibliography = doc.styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
    bibliography.font.name = "Times New Roman"
    force_font_family(bibliography._element.rPr.rFonts, "Times New Roman")
    bibliography.font.size = Pt(12)
    bibliography.paragraph_format.first_line_indent = Cm(-0.75)
    bibliography.paragraph_format.left_indent = Cm(0.75)
    bibliography.paragraph_format.line_spacing = 1.0
    bibliography.paragraph_format.space_after = Pt(3)

    for style in (normal, h1, h2, h3, caption, compact, bibliography):
        set_language(style._element.get_or_add_rPr())

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    add_field(p, "PAGE", "1")

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_body(doc: Document, text: str, *, bold_lead: str | None = None, italic=False) -> None:
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)


def add_noindent(
    doc: Document,
    text: str,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    size=14,
    bold=False,
    line_spacing=1.5,
    space_before=0,
    space_after=0,
    left_indent_cm=None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent_cm is not None:
        p.paragraph_format.left_indent = Cm(left_indent_cm)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_list(doc: Document, items: list[str], numbered=False) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prefix = f"{idx}. " if numbered else "- "
        r = p.add_run(prefix + item)
        set_run_font(r)


def add_equation(doc: Document, expression: str, number: str) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1100, 7438, 1100])
    set_row_cant_split(table.rows[0])
    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            tag = borders.find(qn(f"w:{edge}"))
            if tag is None:
                tag = OxmlElement(f"w:{edge}")
                borders.append(tag)
            tag.set(qn("w:val"), "nil")
    p_mid = table.cell(0, 1).paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.first_line_indent = Cm(0)
    p_mid.paragraph_format.space_before = Pt(6)
    p_mid.paragraph_format.space_after = Pt(6)
    p_mid.paragraph_format.keep_together = True
    run = p_mid.add_run(expression)
    set_run_font(run, size=13, italic=True)
    p_num = table.cell(0, 2).paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.paragraph_format.first_line_indent = Cm(0)
    run = p_num.add_run(f"({number})")
    set_run_font(run)


def add_table(
    doc: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    font_size=10,
    page_break_before=False,
) -> None:
    caption = caption.replace(" - ", " — ")
    font_size = max(font_size, 12)
    cap = doc.add_paragraph()
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_before = Pt(9)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.page_break_before = page_break_before
    run = cap.add_run(caption)
    set_run_font(run, size=12)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    set_row_cant_split(header_row)
    for idx, header in enumerate(headers):
        cell = header_row.cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True)
    for data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(data):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths)


def add_figure(doc: Document, path: Path, caption: str, width_cm=16.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    normalized_caption = caption.replace(" - ", " — ")
    inline_shape = p.add_run().add_picture(str(path), width=Cm(width_cm))
    inline_shape._inline.docPr.set("title", normalized_caption.split(" — ", 1)[0])
    inline_shape._inline.docPr.set("descr", normalized_caption)
    cap = doc.add_paragraph(normalized_caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def load_font(size: int, bold=False) -> ImageFont.FreeTypeFont:
    path = Path("C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf")
    return ImageFont.truetype(str(path), size=size)


def wrap(draw: ImageDraw.ImageDraw, text: str, font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        candidate = (line + " " + word).strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def arrow(draw, start, end, fill=FIG_LINE, width=4, head=14) -> None:
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    p1 = (end[0] - head * math.cos(angle - math.pi / 6), end[1] - head * math.sin(angle - math.pi / 6))
    p2 = (end[0] - head * math.cos(angle + math.pi / 6), end[1] - head * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=fill)


def box(
    draw,
    xy,
    title,
    subtitle="",
    fill=FIG_GRAY_FILL,
    outline=FIG_BLUE,
    title_size=34,
    body_size=27,
) -> None:
    x1, y1, x2, y2 = xy
    draw.rectangle(xy, fill=fill, outline=FIG_INK, width=2)
    draw.line((x1 + 4, y1 + 3, x1 + 4, y2 - 3), fill=outline, width=8)
    title_font = load_font(title_size, bold=True)
    body_font = load_font(body_size)
    title_lines = wrap(draw, title, title_font, x2 - x1 - 30)
    y = y1 + 18
    for line in title_lines:
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=title_font, fill=FIG_INK)
        y += title_size + 5
    if subtitle:
        y += 4
        for line in wrap(draw, subtitle, body_font, x2 - x1 - 30):
            bbox = draw.textbbox((0, 0), line, font=body_font)
            draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=body_font, fill=FIG_INK)
            y += body_size + 4


def canvas(title: str, height=940) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1800, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((60, 38), title, font=load_font(36, bold=True), fill=FIG_INK)
    draw.line((60, 92, 1740, 92), fill=FIG_LINE, width=2)
    return img, draw


def save(img: Image.Image, name: str) -> Path:
    path = FIG_DIR / name
    img.save(path, dpi=(300, 300), optimize=True)
    return path


def make_figures() -> dict[str, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    figs: dict[str, Path] = {}

    img, d = canvas("Этапы развития системы управления")
    stages = [
        ("1. Скалярное V/f", "Реализовано: 0 < f ≤ 50 Гц, MCSDK, безопасный UART"),
        ("2. Векторное FOC", "После идентификации Rs, Rr, Ls, Lr, Lm и J"),
        ("3. ИИ-настройка FOC", "Адаптация параметров и коэффициентов в допустимой области"),
        ("4. ИИ-формирование ШИМ", "Горизонтный выбор векторов с защитным шлюзом"),
    ]
    colors = [
        (FIG_BLUE_FILL, FIG_BLUE),
        (FIG_GRAY_FILL, FIG_LINE),
        (FIG_GRAY_FILL, FIG_LINE),
        (FIG_AMBER_FILL, FIG_AMBER),
    ]
    for i, ((title, subtitle), (fill, outline)) in enumerate(zip(stages, colors)):
        y1 = 150 + i * 180
        box(d, (150, y1, 1650, y1 + 120), title, subtitle, fill=fill, outline=outline)
        if i < 3:
            arrow(d, (900, y1 + 120), (900, y1 + 174), fill=FIG_LINE)
    figs["roadmap"] = save(img, "fig01_roadmap.png")

    img, d = canvas("Функциональная архитектура лабораторного комплекса")
    xs = [70, 370, 680, 990, 1300, 1560]
    specs = [
        ("UNO Q", "Верхний уровень, задания, ИИ"),
        ("ISO7721", "Гальваническая развязка UART"),
        ("NUCLEO G431RB", "MCSDK, V/f и будущий FOC"),
        ("IHM09M2", "Переходник 34 контакта"),
        ("IPM15B", "3-фазный IPM, 125-400 В DC"),
        ("АИР56В2", "0,25 кВт, 220 В Δ"),
    ]
    widths = [250, 230, 270, 260, 230, 190]
    for i, ((title, subtitle), x, w) in enumerate(zip(specs, xs, widths)):
        box(
            d,
            (x, 280, x + w, 520),
            title,
            subtitle,
            fill=FIG_BLUE_FILL if i in (0, 2) else FIG_GRAY_FILL,
            outline=FIG_BLUE if i in (0, 2) else FIG_LINE,
        )
        if i < len(specs) - 1:
            arrow(d, (x + w, 400), (xs[i + 1] - 15, 400))
    d.text((430, 560), "PB7/RX ← UNO D1/TX     PB6/TX → UNO D0/RX     115200 8N1", font=load_font(30), fill=FIG_INK)
    d.text((415, 635), "Break/OCP независимы от UART; аппаратный E-stop ещё требуется", font=load_font(28, bold=True), fill=FIG_INK)
    d.line((410, 674, 1390, 674), fill=FIG_RED, width=3)
    box(d, (1120, 700, 1600, 850), "Измерительные каналы", "Токи фаз, Udc, температура, скорость", fill=FIG_GRAY_FILL, outline=FIG_LINE)
    arrow(d, (1510, 700), (1440, 525), fill=FIG_LINE)
    figs["architecture"] = save(img, "fig02_architecture.png")

    img, d = canvas("Закон скалярного управления V/f и область ослабления поля")
    x0, y0, x1, y1 = 180, 780, 1640, 160
    d.line((x0, y0, x1, y0), fill="#111111", width=4)
    d.line((x0, y0, x0, y1), fill="#111111", width=4)
    arrow(d, (x1 - 20, y0), (x1, y0), fill="#111111", width=4)
    arrow(d, (x0, y1 + 20), (x0, y1), fill="#111111", width=4)
    x50 = x0 + int((50 / 80) * (x1 - x0))
    y220 = y0 - int((220 / 250) * (y0 - y1))
    yboost = y0 - int((20 / 250) * (y0 - y1))
    d.line((x0, yboost, x50, y220), fill=FIG_BLUE, width=6)
    d.line((x50, y220, x1, y220), fill=FIG_RED, width=6)
    d.line((x50, y0, x50, y220), fill=FIG_LINE, width=2)
    d.text((x50 - 65, y0 + 15), "50 Гц", font=load_font(30), fill="#111111")
    d.text((x0 - 135, y220 - 15), "220 В", font=load_font(30), fill="#111111")
    d.text((x0 - 115, yboost - 15), "Uboost", font=load_font(28), fill="#111111")
    d.text((560, 430), "Постоянный поток", font=load_font(36, bold=True), fill=FIG_INK)
    d.text((1210, 430), "Ослабление поля", font=load_font(36, bold=True), fill=FIG_INK)
    d.text((750, 820), "Частота статора f1, Гц", font=load_font(32), fill="#111111")
    d.text((35, 115), "Линейное напряжение U1", font=load_font(32), fill="#111111")
    figs["vf"] = save(img, "fig03_vf_curve.png")

    img, d = canvas("Синхронная и номинальная частота вращения АИР56В2")
    x0, y0, x1, y1 = 190, 790, 1650, 160
    d.line((x0, y0, x1, y0), fill="#111111", width=4)
    d.line((x0, y0, x0, y1), fill="#111111", width=4)
    pts = []
    for f in range(0, 61):
        x = x0 + f / 60 * (x1 - x0)
        n = 60 * f
        y = y0 - n / 3600 * (y0 - y1)
        pts.append((x, y))
    d.line(pts, fill=FIG_BLUE, width=5)
    x50 = x0 + 50 / 60 * (x1 - x0)
    y3000 = y0 - 3000 / 3600 * (y0 - y1)
    y2720 = y0 - 2720 / 3600 * (y0 - y1)
    d.ellipse((x50 - 10, y3000 - 10, x50 + 10, y3000 + 10), fill="white", outline=FIG_BLUE, width=4)
    d.rectangle((x50 - 9, y2720 - 9, x50 + 9, y2720 + 9), fill=FIG_RED)
    d.line((x50, y2720, x50, y3000), fill=FIG_RED, width=3)
    d.text((x50 - 80, y3000 - 55), "n_s = 3000 об/мин", font=load_font(28), fill=FIG_INK)
    d.text((x50 - 80, y2720 + 15), "n_N = 2720 об/мин", font=load_font(28), fill=FIG_INK)
    d.text((1000, 640), "Номинальное скольжение s = 0,093", font=load_font(32, bold=True), fill=FIG_INK)
    d.text((760, 825), "Частота f1, Гц", font=load_font(32), fill="#111111")
    d.text((25, 115), "Частота вращения, об/мин", font=load_font(32), fill="#111111")
    figs["speed"] = save(img, "fig04_speed.png")

    img, d = canvas("Структура векторного управления с ориентацией по потоку ротора")
    blocks = [
        (80, 300, 330, 520, "Задания", "n*, ψ*"),
        (410, 300, 700, 520, "Регуляторы", "PI скорости, id, iq"),
        (790, 300, 1040, 520, "dq → αβ", "обратный Park"),
        (1130, 300, 1370, 520, "SVPWM", "Ta, Tb, Tc"),
        (1460, 300, 1710, 520, "Инвертор + АД", "U, V, W"),
    ]
    for i, (x1b, y1b, x2b, y2b, t, s) in enumerate(blocks):
        box(
            d,
            (x1b, y1b, x2b, y2b),
            t,
            s,
            fill=FIG_BLUE_FILL if i < 3 else FIG_TEAL_FILL,
            outline=FIG_BLUE if i < 3 else FIG_TEAL,
        )
        if i < len(blocks) - 1:
            arrow(d, (x2b, 410), (blocks[i + 1][0] - 15, 410))
    box(d, (790, 670, 1120, 850), "Наблюдатель", "θe, ψr, ω", fill=FIG_AMBER_FILL, outline=FIG_AMBER)
    box(d, (380, 670, 690, 850), "Clarke/Park", "ia, ib, ic → id, iq", fill=FIG_GRAY_FILL, outline=FIG_LINE)
    arrow(d, (1580, 525), (1580, 755), fill=FIG_LINE)
    arrow(d, (1580, 755), (1120, 755), fill=FIG_LINE)
    arrow(d, (790, 755), (690, 755), fill=FIG_LINE)
    arrow(d, (535, 670), (535, 525), fill=FIG_LINE)
    figs["foc"] = save(img, "fig05_foc.png")

    img, d = canvas("Векторная диаграмма двухуровневого инвертора")
    cx, cy, radius = 900, 500, 310
    points = []
    for k in range(6):
        a = -math.pi / 2 + k * math.pi / 3
        points.append((cx + radius * math.cos(a), cy + radius * math.sin(a)))
    d.polygon(points, outline=FIG_BLUE, width=4)
    for k, pt in enumerate(points, start=1):
        arrow(d, (cx, cy), pt, fill=FIG_TEAL, width=4)
        d.text((pt[0] - 28, pt[1] - 28), f"V{k}", font=load_font(28, bold=True), fill=FIG_INK)
    arrow(d, (cx, cy), (cx + 205, cy - 125), fill=FIG_RED, width=6)
    d.text((cx + 220, cy - 170), "V*", font=load_font(34, bold=True), fill=FIG_INK)
    d.ellipse((cx - 10, cy - 10, cx + 10, cy + 10), fill="#111111")
    d.text((120, 740), "Допустимое действие выбирается из конечного множества состояний ключей;", font=load_font(31), fill="#111111")
    d.text((120, 790), "защитный шлюз исключает сквозной ток, короткие импульсы и нарушение dead-time.", font=load_font(31), fill="#111111")
    figs["svpwm"] = save(img, "fig06_svpwm.png")

    img, d = canvas("Формат командного кадра UNO Q — Nucleo, версия 0x02")
    fields = [
        ("AA 55", 2, "Синхронизация", BLUE),
        ("02", 1, "Версия", TEAL),
        ("flags", 1, "START/CLEAR/ESTOP", AMBER),
        ("mode", 1, "0 или 3", BLUE),
        ("seq", 1, "Счётчик", GRAY),
        ("f_mHz", 4, "Частота LE", TEAL),
        ("legacy", 4, "Не используется", GRAY),
        ("0...0", 17, "Резерв = 0", LIGHT_GRAY),
        ("CRC", 1, "XOR 0..30", RED),
    ]
    x = 70
    total = 32
    for label, count, desc, color in fields:
        w = max(70, int(1640 * count / total))
        field_fills = {
            BLUE: FIG_BLUE_FILL,
            TEAL: FIG_TEAL_FILL,
            AMBER: FIG_AMBER_FILL,
            RED: FIG_RED_FILL,
            GRAY: FIG_GRAY_FILL,
            LIGHT_GRAY: "#FAFAFA",
        }
        fill = field_fills.get(color, FIG_GRAY_FILL)
        d.rectangle((x, 250, x + w, 460), fill=fill, outline="#111111", width=3)
        text_color = FIG_INK
        for j, line in enumerate(wrap(d, label, load_font(28, bold=True), w - 12)):
            d.text((x + 7, 275 + j * 34), line, font=load_font(28, bold=True), fill=text_color)
        for j, line in enumerate(wrap(d, desc, load_font(22), w - 12)):
            d.text((x + 7, 365 + j * 27), line, font=load_font(22), fill=text_color)
        d.text((x + 7, 475), f"{count} байт", font=load_font(21), fill="#111111")
        x += w
    box(d, (190, 620, 720, 820), "Принимается", "mode=3; 0 < f ≤ 50 Гц; резерв нулевой; CRC верен", fill=FIG_TEAL_FILL, outline=FIG_TEAL)
    box(d, (1070, 620, 1600, 820), "Отклоняется", "FOC/DUTY/DIAG; f > 50 Гц; timeout > 300 мс", fill=FIG_RED_FILL, outline=FIG_RED)
    arrow(d, (720, 720), (1060, 720), fill=FIG_LINE)
    figs["uart"] = save(img, "fig07_uart.png")

    img, d = canvas("Состояния безопасного запуска и остановки привода")
    nodes = {
        "SAFE": (160, 330, 430, 520, "SAFE", "PWM запрещён"),
        "READY": (570, 330, 840, 520, "READY", "Требование: связь и защиты"),
        "RUN": (980, 330, 1250, 520, "RUN", "V/f, 0 < f ≤ 50 Гц"),
        "FAULT": (1390, 330, 1660, 520, "FAULT", "Защёлкнутый отказ"),
        "ESTOP": (760, 670, 1040, 850, "E-STOP", "Отдельная цепь требуется"),
    }
    for name, (x1b, y1b, x2b, y2b, t, s) in nodes.items():
        fill, outline = (FIG_RED_FILL, FIG_RED) if name in ("FAULT", "ESTOP") else (FIG_BLUE_FILL, FIG_BLUE)
        box(d, (x1b, y1b, x2b, y2b), t, s, fill=fill, outline=outline)
    arrow(d, (430, 425), (560, 425), fill=FIG_TEAL)
    arrow(d, (840, 425), (970, 425), fill=FIG_TEAL)
    arrow(d, (1250, 425), (1380, 425), fill=FIG_RED)
    arrow(d, (1115, 520), (970, 660), fill=FIG_RED)
    arrow(d, (760, 755), (430, 520), fill=FIG_RED)
    d.text((455, 370), "CLEAR", font=load_font(25), fill=FIG_INK)
    d.text((865, 370), "START", font=load_font(25), fill=FIG_INK)
    d.text((1260, 370), "fault/timeout", font=load_font(25), fill=FIG_INK)
    figs["states"] = save(img, "fig08_states.png")

    img, d = canvas("Предлагаемая иерархия интеллектуального управления")
    box(d, (120, 160, 1680, 300), "Уровень 3: UNO Q / вычислительный узел", "Обучение, планирование эксперимента, выбор ограниченных параметров", fill=FIG_AMBER_FILL, outline=FIG_AMBER)
    box(d, (250, 385, 1550, 545), "Уровень 2: адаптация", "Оценка Rs, Rr, Lm; настройка PI; весовые коэффициенты; доверие модели", fill=FIG_TEAL_FILL, outline=FIG_TEAL)
    box(d, (380, 635, 1420, 800), "Уровень 1: детерминированное реальное время", "FOC + SVPWM + break + OCP + watchdog + Safety Gateway", fill=FIG_BLUE_FILL, outline=FIG_BLUE)
    arrow(d, (900, 300), (900, 375), fill=FIG_AMBER)
    arrow(d, (900, 545), (900, 625), fill=FIG_TEAL)
    d.text((1080, 330), "только ограниченные команды", font=load_font(25), fill=FIG_INK)
    d.text((1040, 575), "проекция на допустимое множество", font=load_font(25), fill=FIG_INK)
    figs["hierarchy"] = save(img, "fig09_hierarchy.png")

    img, d = canvas("Контур обучения и верификации интеллектуального регулятора")
    blocks = [
        (80, 260, 340, 470, "Паспорт и измерения", "Rs, Rr, Lm, J, B"),
        (430, 260, 690, 470, "Цифровой двойник", "dq-модель + инвертор"),
        (780, 260, 1040, 470, "Обучение", "domain randomization"),
        (1130, 260, 1390, 470, "Safety Gateway", "ограничения и fallback"),
        (1480, 260, 1720, 470, "HIL / стенд", "A/B испытания"),
    ]
    for i, item in enumerate(blocks):
        box(
            d,
            item[:4],
            item[4],
            item[5],
            fill=FIG_BLUE_FILL if i % 2 == 0 else FIG_TEAL_FILL,
            outline=FIG_BLUE if i % 2 == 0 else FIG_TEAL,
        )
        if i < len(blocks) - 1:
            arrow(d, (item[2], 365), (blocks[i + 1][0] - 15, 365))
    box(d, (620, 650, 1180, 835), "Реестр доказательств", "хэши, сценарии, метрики, версии, отрицательные результаты", fill=FIG_GRAY_FILL, outline=FIG_LINE)
    arrow(d, (1600, 470), (1180, 735), fill=FIG_LINE)
    arrow(d, (620, 735), (210, 470), fill=FIG_LINE)
    figs["ai_pipeline"] = save(img, "fig10_ai_pipeline.png")

    img, d = canvas("Последовательность безопасных экспериментальных работ")
    steps = [
        ("S0", "Статическая проверка", "сборка, CRC, ограничения"),
        ("S1", "Низковольтный bench", "без J7 и без двигателя"),
        ("S2", "HIL / имитатор", "fault, timeout, E-stop"),
        ("S3", "V/f без нагрузки", "ограниченный источник"),
        ("S4", "FOC", "после идентификации"),
        ("S5", "ИИ A/B", "только после базовых тестов"),
    ]
    for i, (code, title, sub) in enumerate(steps):
        x = 60 + i * 285
        fill = FIG_TEAL_FILL if i == 0 else FIG_GRAY_FILL
        outline = FIG_TEAL if i == 0 else FIG_LINE
        box(d, (x, 280, x + 230, 540), f"{code}: {title}", sub, fill=fill, outline=outline, title_size=28, body_size=24)
        if i < len(steps) - 1:
            arrow(d, (x + 230, 410), (x + 270, 410), fill=FIG_LINE)
    d.text((430, 665), "Переход на следующий этап допускается только после выполнения критериев предыдущего", font=load_font(31, bold=True), fill=FIG_INK)
    d.line((420, 712, 1450, 712), fill=FIG_RED, width=3)
    figs["test_steps"] = save(img, "fig11_test_steps.png")

    img, d = canvas("Границы подтверждённых результатов рабочей редакции")
    box(d, (120, 180, 790, 780), "Подтверждено", "Числовая согласованность профиля AIR56B2\nMCSDK: успешная сборка\nELF, BIN, HEX и SHA-256\nСтатический контракт UART v0x02\nОграничение 50 Гц\n21 текущий host-тест", fill=FIG_TEAL_FILL, outline=FIG_TEAL)
    box(d, (1010, 180, 1680, 780), "Не подтверждено", "Шильдик конкретного экземпляра\nПолная схема Nucleo/IHM/IPM\nАппаратный E-stop, предзаряд и HIL\nПараметры Rs, Rr, Lls, Llr, Lm, J\nFOC и тяговый масштаб\nЭкспериментальное преимущество ИИ", fill=FIG_RED_FILL, outline=FIG_RED)
    arrow(d, (790, 480), (995, 480), fill=FIG_LINE)
    d.text((815, 405), "экспериментальная", font=load_font(24), fill=FIG_INK)
    d.text((845, 440), "программа", font=load_font(24), fill=FIG_INK)
    figs["evidence"] = save(img, "fig12_evidence.png")

    return figs


def add_title_page(doc: Document) -> None:
    add_noindent(
        doc,
        "ФЕДЕРАЛЬНОЕ АГЕНТСТВО ЖЕЛЕЗНОДОРОЖНОГО ТРАНСПОРТА",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        line_spacing=1.0,
        space_after=3,
    )
    add_noindent(
        doc,
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        line_spacing=1.0,
        space_after=3,
    )
    add_noindent(
        doc,
        "«ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ ПУТЕЙ СООБЩЕНИЯ ИМПЕРАТОРА АЛЕКСАНДРА I»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=12,
        bold=True,
        line_spacing=1.0,
    )
    add_noindent(
        doc,
        "На правах рукописи",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=12,
        line_spacing=1.0,
        space_before=24,
    )
    add_noindent(
        doc,
        "[ФАМИЛИЯ ИМЯ ОТЧЕСТВО]",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        bold=True,
        line_spacing=1.0,
        space_before=30,
    )
    add_noindent(
        doc,
        "ПОВЫШЕНИЕ ЭНЕРГЕТИЧЕСКОЙ ЭФФЕКТИВНОСТИ АСИНХРОННОГО ТЯГОВОГО ЭЛЕКТРОПРИВОДА НА ОСНОВЕ ИНТЕЛЛЕКТУАЛЬНОЙ АДАПТАЦИИ АЛГОРИТМОВ ВЕКТОРНОГО УПРАВЛЕНИЯ",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        bold=True,
        line_spacing=1.0,
        space_before=24,
    )
    add_noindent(
        doc,
        "Специальность 2.9.3 — Подвижной состав железных дорог, тяга поездов и электрификация",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        line_spacing=1.0,
        space_before=24,
    )
    add_noindent(
        doc,
        "Диссертация на соискание учёной степени кандидата технических наук",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        line_spacing=1.0,
        space_before=12,
    )
    add_noindent(
        doc,
        "Научный руководитель:\n[Ф.И.О., учёная степень, учёное звание]",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        size=12,
        line_spacing=1.0,
        space_before=120,
        left_indent_cm=9.0,
    )
    add_noindent(
        doc,
        "Санкт-Петербург — 2026",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=14,
        line_spacing=1.0,
        space_before=230,
    )
    doc.add_page_break()


def add_status_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.page_break_before = False
    r = p.add_run("СТАТУС РУКОПИСИ")
    set_run_font(r, bold=True)
    add_body(doc, "Настоящий документ является развёрнутой рабочей редакцией диссертации, подготовленной на основе фактических исходных кодов, конфигураций, каталожного профиля двигателя и результатов программной проверки проекта MIC_AI. Поля с фамилией соискателя, научным руководителем, апробацией, публикациями и актами внедрения оставлены для заполнения после согласования с научным руководителем.")
    add_body(doc, "В документе строго разделены три класса доказательств: подтверждённые сборкой и статическими проверками результаты; результаты моделирования на персональном компьютере; будущие стендовые и высоковольтные испытания. Моделирование не выдается за физический эксперимент, а подготовленный образ V/f не считается допуском к силовому пуску до измерения параметров двигателя, проверки предзаряда и аппаратных защит.")
    add_table(doc, "Таблица 1 - Статус основных частей исследования", ["Часть", "Статус", "Разрешённая формулировка"], [
        ["Скалярное V/f", "Программно реализовано", "Прошивка-кандидат; требуется стендовая проверка"],
        ["Векторное FOC", "Методика и модель", "Реализация после идентификации двигателя"],
        ["ИИ-настройка", "Предложенная архитектура", "Host-моделирование; без аппаратного доказательства"],
        ["ИИ-ШИМ", "Исследовательский прототип", "Без утверждения о превосходстве и готовности к MCU"],
    ], [2500, 2200, 4938], font_size=10)
    add_body(doc, "Перед подачей в диссертационный совет рабочая редакция должна быть дополнена протоколами экспериментов, подтверждёнными публикациями по теме, сведениями об апробации, актами внедрения при их наличии и итоговыми выводами, согласованными с научным руководителем.")
    doc.add_page_break()


def add_toc_entry(doc: Document, title: str, page: str, level: int = 0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.65 * level)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    r = p.add_run(f"{title}\t{page}")
    set_run_font(r, size=12, bold=(level == 0))


def add_toc_and_abbreviations(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run("ОГЛАВЛЕНИЕ")
    set_run_font(r, bold=True)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    toc.paragraph_format.line_spacing = 1.0
    add_field(toc, 'TOC \\o "1-2" \\h \\z \\u', "Обновите оглавление: Ctrl+A, F9")
    doc.add_page_break()

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.paragraph_format.page_break_before = False
    r = p.add_run("СПИСОК СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ")
    set_run_font(r, bold=True)
    abbrev = [
        ("АД", "асинхронный двигатель"),
        ("ВАК", "Высшая аттестационная комиссия при Минобрнауки России"),
        ("ДПТ", "дискретное преобразование координат и/или данных"),
        ("ИИ", "искусственный интеллект"),
        ("КПД", "коэффициент полезного действия"),
        ("МК", "микроконтроллер"),
        ("ПЧ", "преобразователь частоты"),
        ("ШИМ", "широтно-импульсная модуляция"),
        ("ACIM", "asynchronous induction motor"),
        ("FOC", "field-oriented control, векторное управление с ориентацией по полю"),
        ("HIL", "hardware-in-the-loop, испытание аппаратуры в замкнутом контуре"),
        ("MCSDK", "STM32 Motor Control Software Development Kit"),
        ("OCP", "аппаратная защита от сверхтока"),
        ("PI", "пропорционально-интегральный регулятор"),
        ("SVPWM", "space-vector pulse-width modulation"),
        ("UART", "асинхронный последовательный интерфейс"),
        ("V/f", "скалярное управление с заданным отношением напряжения к частоте"),
    ]
    add_table(doc, "Таблица 2 - Принятые сокращения", ["Обозначение", "Расшифровка"], [[a, b] for a, b in abbrev], [2200, 7438], font_size=11)


def add_intro(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("ВВЕДЕНИЕ", level=1)
    paragraphs = [
        "Актуальность темы исследования определяется необходимостью повышения энергетической эффективности, управляемости и диагностируемости тягового электропривода железнодорожного подвижного состава. Асинхронная машина в сочетании с полупроводниковым преобразователем обладает высокой конструктивной надёжностью, однако фактические потери, динамические качества и устойчивость управления зависят от качества идентификации параметров, структуры регуляторов, закона модуляции и ограничений силовой части. Для транспортного применения особенно важны воспроизводимость алгоритмов, отказобезопасное состояние и возможность доказать границы применимости результата.",
        "Паспорт научной специальности 2.9.3 включает исследования тягового привода и силовых преобразователей, автоматизацию и управление подвижным составом, оценку энергетических потерь, а также компьютерное моделирование и испытания. Разрабатываемый подход соответствует этим направлениям, поскольку связывает модель асинхронного двигателя, микроконтроллерное управление преобразователем, программно-аппаратные защиты и методику сравнительных испытаний [1]. Лабораторный двигатель мощностью 0,25 кВт рассматривается как масштабная исследовательская установка, а не как непосредственный тяговый двигатель.",
        "Степень разработанности темы характеризуется наличием хорошо исследованных законов V/f, векторного управления, прямого управления моментом и прогнозирующего управления [6-12]. Современные работы расширяют классические контуры нейросетевой идентификацией, настройкой регуляторов и оцениванием переменных состояния [13-18]. Вместе с тем публикационная новизна не может сводиться к простой замене PI-регулятора нейронной сетью. Требуется архитектура, в которой обучаемая часть действует в формально ограниченной области, а детерминированный контур реального времени сохраняет приоритет защиты.",
        "Исходная инженерная база исследования включает NUCLEO-G431RB на микроконтроллере семейства STM32G431RBTx, переходную плату X-NUCLEO-IHM09M2, силовой модуль STEVAL-IPM15B, асинхронный двигатель IEK АИР56В2 и вычислительный узел Arduino UNO Q. Для этой конфигурации подготовлена сборка официального примера MCSDK ACIM V/f Open Loop, разработан программный контракт изолированного интерфейса команд и сформирован комплект ELF, BIN и HEX. Полная электрическая схема текущей конфигурации, аппаратный UART-тракт, силовые платы и двигатель ещё не прошли совместные стендовые испытания, поэтому текущий результат является программно проверенным кандидатом.",
    ]
    for t in paragraphs:
        add_body(doc, t)
    add_body(doc, "Структура и оформление рабочей редакции приняты с учётом ГОСТ Р 7.0.11-2011 [2], действующего порядка присуждения учёных степеней [3], регламента диссертационного совета ПГУПС [4] и доступного примера защищённой диссертации по тематике подвижного состава [5]. Служебная страница статуса сохраняется только в рабочей редакции и перед официальной подачей удаляется.")

    add_body(doc, "Цель работы — разработать и экспериментально проверить метод ограниченной адаптации векторного управления, обеспечивающий статистически подтверждённое снижение удельных энергозатрат относительно настроенного FOC-SVPWM при неухудшении заданных показателей динамики и безопасности в установленном множестве параметрических и режимных возмущений.", bold_lead="Цель работы")
    add_body(doc, "Для достижения цели поставлены следующие задачи:")
    add_list(doc, [
        "выполнить анализ методов управления асинхронным тяговым электроприводом и определить границы применения скалярного, векторного и интеллектуального управления;",
        "разработать математическую модель асинхронного двигателя, преобразователя и измерительного тракта, пригодную для идентификации и сравнительного моделирования;",
        "подготовить воспроизводимую программную конфигурацию и проект подключения платформы на базе NUCLEO-G431RB, STEVAL-IPM15B, АИР56В2 и UNO Q;",
        "реализовать безопасный базовый режим V/f как исходную точку стендовой проверки;",
        "разработать метод векторного управления и процедуру идентификации параметров целевого двигателя;",
        "разработать метод ограниченной адаптации параметров модели и регуляторов FOC с детерминированной проверкой допустимости каждой коррекции;",
        "сформировать программу HIL- и стендовых испытаний, систему метрик и правила статистического сравнения;",
        "проверить программные контракты, воспроизводимость сборки и прослеживаемость результатов."], numbered=True)

    add_body(doc, "Объект исследования — процессы электромеханического преобразования энергии и управления в асинхронном тяговом электроприводе с автономным инвертором напряжения.", bold_lead="Объект исследования")
    add_body(doc, "Предмет исследования — закономерности влияния параметрической неопределённости и ограниченной интеллектуальной адаптации модели и регуляторов FOC на энергетические, динамические и защитные показатели электропривода.", bold_lead="Предмет исследования")
    add_body(doc, "Методы исследования включают теорию электрических машин и электропривода, преобразования координат, теорию автоматического управления, численное моделирование, идентификацию параметров, анализ конечного множества состояний инвертора, методы машинного обучения, статическую проверку исходного кода и экспериментальный дизайн.", bold_lead="Методы исследования")

    add_body(doc, "Научная гипотеза состоит в том, что энергетические и динамические показатели асинхронного привода могут быть улучшены без передачи обучаемому алгоритму неограниченного доступа к ключам инвертора, если разделить систему на детерминированный контур FOC, медленный адаптивный уровень и Safety Gateway, проектирующий каждое интеллектуальное действие на допустимое множество.", bold_lead="Научная гипотеза")
    add_body(doc, "К проверяемым элементам научной новизны относятся: метод событийной ограниченной адаптации параметров модели и регуляторов FOC, отличающийся проекцией обновлений на физически допустимое множество и блокировкой адаптации при снижении достоверности идентификации; критерий допустимости интеллектуальной коррекции, совместно учитывающий ошибку регулирования, энергетические потери, коммутационную нагрузку и запас до ограничений; процедура доказательного сравнения адаптивного и базового управления при одинаковых режимах и ограничениях. Окончательные формулировки должны содержать численный эффект, доверительный интервал и область применимости.", bold_lead="Проверяемые элементы научной новизны")
    add_body(doc, "Теоретическая значимость работы состоит в формализации условий включения ограниченного адаптивного уровня в детерминированную структуру FOC, определении допустимого множества изменяемых параметров и критериев перехода к базовому управлению при нарушении ограничений, выходе за область идентификации или превышении вычислительного времени.", bold_lead="Теоретическая значимость")
    add_body(doc, "Практическая значимость состоит в создании платформы, где версия протокола, профиль двигателя и хэши бинарных артефактов фиксируются совместно. Текущий SHA-256 однозначно идентифицирует готовый файл, но полная воспроизводимость происхождения потребует дополнительно хэшировать профиль, исходники, Git commit, версию компилятора и команду сборки. Такой подход снижает риск случайной прошивки неподходящего образа и позволяет постепенно переходить от V/f к FOC и интеллектуальному управлению.", bold_lead="Практическая значимость")
    add_body(doc, "Положения, подлежащие экспериментальной проверке, в рабочей редакции сформулированы как проверяемые утверждения; численные поля заполняются только после стендовой серии:")
    add_list(doc, [
        "проекция событийного обновления на Θ сохраняет θk∈Θ при θ0∈Θ, а блокировка по информативности, доверию и времени выдержки предотвращает изменение оценки вне разрешённых условий; влияние на удельную энергию составляет [ΔE, CI95] в области [режимы];",
        "логически обособленная детерминированная проверка коррекции исключает её применение при нарушении токовых, напряженческих, тепловых и временных ограничений; полнота обнаружения составляет [результат fault-injection];",
        "прослеживаемая цепочка профиль—исходники—сборка—манифест—протокол обеспечивает воспроизводимость программной конфигурации с точностью до зафиксированных хэшей и версий;",
        "перенос лабораторного результата к тяговой задаче допустим только через модель (2.18)-(2.19), тяговый цикл и ограничения сцепления; численная применимость устанавливается для [серия подвижного состава, цикл]."], numbered=True)
    add_body(doc, "Достоверность будущих выводов будет обеспечиваться использованием официальной документации компонентов, независимой проверкой программных контрактов, сопоставлением расчётной и измеренной модели, повторными опытами, доверительными интервалами, A/B-сравнением на одинаковых режимах и публикацией отрицательных результатов. На текущем этапе достоверно подтверждены только перечисленные в главе 5 программные проверки.", bold_lead="Достоверность результатов")
    add_body(doc, "Апробация и публикации: [заполнить фактическими докладами, статьями из перечня ВАК, свидетельствами о регистрации программ и иными результатами после их появления].")
    add_body(doc, "Структура работы. Диссертация включает введение, пять глав, заключение, список литературы и приложения. В первой главе сформулирована научная задача; во второй приведены модели; в третьей описана программно-аппаратная платформа; в четвёртой предложен интеллектуальный метод; в пятой представлены подтверждённые программные результаты и программа эксперимента.")
    add_figure(doc, figs["roadmap"], "Рисунок 1 - Логика поэтапного развития исследуемой системы управления")
    add_body(doc, "Последовательность на рисунке 1 является обязательной логикой получения доказательств. Каждый следующий уровень наследует измеренные ограничения предыдущего: FOC вводится после безопасной проверки силового тракта в режиме V/f; интеллектуальная настройка проверяется относительно устойчивого FOC; интеллектуальное формирование ШИМ допускается только после подтверждения логически обособленного защитного шлюза и аппаратных защит. Поэтому незавершённый аппаратный этап не маскируется результатами компьютерного моделирования следующего уровня.")
    add_body(doc, "Рабочая редакция фиксирует исходное состояние исследования и одновременно задаёт проверяемую программу его завершения. После получения стенда расчётные разделы должны быть дополнены протоколами идентификации, осциллограммами, балансом погрешностей и статистикой повторных опытов. Такая организация позволяет обновлять выводы без изменения заранее заявленных критериев качества и исключает выбор метрик после ознакомления с результатом.")


def add_chapter1(doc: Document) -> None:
    doc.add_heading("1 АНАЛИЗ МЕТОДОВ УПРАВЛЕНИЯ АСИНХРОННЫМ ТЯГОВЫМ ЭЛЕКТРОПРИВОДОМ", level=1)
    doc.add_heading("1.1 Требования тягового применения", level=2)
    for t in [
        "Тяговый электропривод работает в широком диапазоне скоростей и нагрузок, включая пуск, ограничение сцепления, ослабление поля, выбег, электрическое торможение и воздействие возмущений со стороны контактной сети. Критерий качества не может ограничиваться ошибкой скорости: существенны удельное энергопотребление, потери в меди и стали, пульсации момента, температура полупроводников, число коммутаций, устойчивость к изменению сопротивлений и поведение при отказах датчиков.",
        "Для железнодорожного подвижного состава отказ алгоритма управления может перейти в опасное состояние силовой части. Поэтому обучаемый компонент допустим только при наличии детерминированного ограничения тока, напряжения, температуры, dead-time и последовательности включения. Связь с верхним уровнем не должна быть единственным каналом аварийной остановки. Эти требования определяют архитектуру, в которой UNO Q задаёт ограниченные параметры, а STM32G431 выполняет жёсткий цикл управления и аппаратные защиты.",
        "Лабораторная установка мощностью 0,25 кВт не воспроизводит электромагнитные, тепловые и механические постоянные тягового двигателя в абсолютных единицах. Научно корректное масштабирование должно использовать относительные величины: ток и поток в долях номинальных, электромагнитный момент в долях базового, частоту коммутации относительно электрической постоянной времени, а тепловую нагрузку относительно допустимого нагрева. Вывод о применимости к тяге требует дополнительной проверки на модели соответствующего класса мощности.",
    ]:
        add_body(doc, t)
    add_body(doc, "Группы показателей и способы их подтверждения сведены в таблицу 1.1.")
    add_table(doc, "Таблица 1.1 - Группы требований к системе управления", ["Группа", "Контролируемые показатели", "Средство подтверждения"], [
        ["Энергетика", "КПД, потери, Udc, ток", "Баланс мощности, анализатор"],
        ["Динамика", "Ошибка скорости, момент, время перехода", "Осциллограммы и идентичные задания"],
        ["Качество тока", "RMS, спектр, THD", "Синхронная выборка токов"],
        ["Безопасность", "OCP, break, E-stop, watchdog", "Fault injection, HIL"],
        ["Воспроизводимость", "Версии, хэши, профиль", "Автоматизированная сборка"],
    ], [1700, 4200, 3738], font_size=10)

    doc.add_heading("1.2 Скалярное управление V/f", level=2)
    for t in [
        "Скалярный закон V/f поддерживает приблизительное постоянство магнитного потока за счёт пропорционального изменения амплитуды основной гармоники напряжения при изменении частоты. Его достоинствами являются простота, малое число требуемых параметров и удобство первичного запуска. Именно поэтому V/f выбран первым этапом: он позволяет проверить фазировку, направления, защиту, связь и силовой тракт до ввода чувствительной к параметрам модели FOC.",
        "Недостатки V/f проявляются при низкой скорости, резком изменении нагрузки и необходимости регулировать момент. Падение напряжения на сопротивлении статора и инверторе нарушает пропорциональность потока, поэтому вводится низкочастотная добавка Uboost. Без обратной связи скольжение определяется нагрузкой и температура ротора влияет на рабочую точку. Следовательно, V/f является базовой линией сравнения, но не конечным способом высокодинамичного тягового управления.",
        "В текущем проекте команда START задаёт частоту статора в диапазоне 0 < f ≤ 50 Гц. Нулевая частота реализуется командами STOP/OFF и не является допустимым заданием запуска. Поле амплитуды в UART-кадре сохранено для совместимости, но Nucleo его игнорирует: профиль V/f формируется средствами MCSDK. Это исключает возможность обхода паспорта двигателя произвольным коэффициентом, поступившим с верхнего уровня.",
    ]:
        add_body(doc, t)

    doc.add_heading("1.3 Векторное управление и идентификация", level=2)
    for t in [
        "Векторное управление преобразует измеренные фазные токи в систему координат, связанную с оцененным потокосцеплением. Компонента тока по оси d регулирует поток, а компонента по оси q — электромагнитный момент. При корректной ориентации достигается приближённое развязывание каналов, аналогичное независимому управлению возбуждением и моментом машины постоянного тока [7-10].",
        "Качество FOC зависит от сопротивлений статора и ротора, взаимной и рассеяющих индуктивностей, параметров датчиков и задержек вычисления. Сопротивление ротора заметно меняется при нагреве, а неточная взаимная индуктивность искажает оценку потока. Для предварительной V/f-конфигурации принятых каталожных данных достаточно, но для FOC их недостаточно. Параметры Rs, Rr, Lls, Llr, Lm и момент инерции должны быть измерены и помещены в версионируемый профиль.",
        "Идентификация должна выполняться до высоковольтного пуска или в безопасной процедуре с ограниченным напряжением и током. Результат следует проверять не только по совпадению одного переходного процесса, но и на нескольких частотах и уровнях нагрузки. Неопределённость параметров становится входом для последующей интеллектуальной адаптации, а не скрывается подбором коэффициентов.",
    ]:
        add_body(doc, t)

    doc.add_heading("1.4 Интеллектуальные методы и требования к доказательству", level=2)
    for t in [
        "Нейросетевые и адаптивные методы применяются для оценивания скорости, идентификации параметров, настройки регуляторов и непосредственного выбора управляющего воздействия [13-17]. Их потенциальное преимущество связано со способностью аппроксимировать нелинейные зависимости и учитывать контекст режима. Одновременно возникают риски поведения вне обучающей выборки, вычислительной недетерминированности и отсутствия физической интерпретируемости.",
        "Научное сравнение должно включать сильные классические базовые алгоритмы: настроенный FOC-SVPWM, DTC, DTC-SVM, одношаговый FCS-MPC и регулятор тока конечного времени. Сравнение с заведомо слабой реализацией не подтверждает новизну. Для каждого алгоритма фиксируются ограничения тока и напряжения, частота дискретизации, вычислительный бюджет и процедура настройки.",
        "В проекте имеется архивный host-прототип Safe Neural Horizon PWM. Он объединяет модельный прогноз, нейросетевую коррекцию стоимости, конечный поиск векторов и защитный шлюз. Прототип прошёл программную матрицу из 31 сценария, но не переносился на MCU, HIL или физический стенд. Поэтому его данные используются для формирования гипотез и программы исследования, а не как доказательство превосходства.",
    ]:
        add_body(doc, t)
    add_body(doc, "Сопоставление основных групп методов, доступных доказательств и ограничений приведено в таблице 1.2. Оно показывает, что исследовательский разрыв связан не с отсутствием ещё одного регулятора, а с недостатком прослеживаемых ограничений, тяговых сценариев и статистически доказанного выигрыша адаптации.")
    add_table(doc, "Таблица 1.2 - Критический обзор методов управления и адаптации", ["Подход", "Подтверждённые достоинства", "Ограничение для настоящей работы"], [
        ["V/f [6, 7]", "Простота, малая параметрическая чувствительность", "Низкая точность момента и низкоскоростных режимов"],
        ["FOC [8-11]", "Раздельное управление потоком и моментом", "Зависимость от Rs, Rr, Lm и оценки состояния"],
        ["DTC, FCS-MPC [12]", "Быстрый отклик и явный выбор состояния", "Пульсации, вычислительная стоимость, настройка стоимости"],
        ["Нейросетевая идентификация [13, 14, 16, 17]", "Аппроксимация нелинейностей и дрейфа", "OOD-риск и недостаток детерминированных ограничений"],
        ["Архивный SNH-PWM", "Воспроизводимый host-прототип и fault-сценарии", "Нет MCU, HIL и физической верификации"],
    ], [2200, 3500, 3938], font_size=9)

    doc.add_heading("1.5 Постановка научной задачи", level=2)
    add_body(doc, "Требуется разработать и экспериментально проверить метод, который сохраняет детерминированную основу FOC, допускает адаптацию к нагреву и неопределённости параметров и обеспечивает статистически подтверждаемое снижение удельных энергозатрат при неухудшении заданных показателей динамики и безопасности в установленном множестве возмущений. Интеллектуальный уровень не должен формировать необработанные сигналы ключей, пока действие не прошло детерминированную проверку допустимости.")
    add_equation(doc, "min J = wₑEloss + wₙeₙ² + wᵢ||eᵢ||² + wₜrₜ + wₛwNₛw + wₛPₛ", "1.1")
    add_body(doc, "Здесь Eloss — энергия потерь или её заранее определённый измеримый прокси, eₙ — ошибка скорости, eᵢ — ошибка токов, rₜ — показатель пульсации момента, Nₛw — число переключений, Pₛ — штраф за приближение к ограничениям. Робастность определяется как сохранение ограничений и допустимого качества во всём заданном множестве вариаций Rs, Rr, Lm, нагрузки, Udc и измерительных ошибок. Весовые коэффициенты и множество возмущений фиксируются до испытаний.")
    doc.add_heading("1.6 Выводы по главе 1", level=2)
    add_body(doc, "Выполненный анализ позволяет сформулировать следующие результаты первой главы:")
    add_list(doc, [
        "скалярный режим V/f принят как обязательная базовая линия, обеспечивающая проверку фазировки, измерительных каналов и защит при минимальном числе неизвестных параметров;",
        "векторное управление выбрано детерминированной основой перспективной системы, поскольку оно разделяет управление потокосцеплением и моментом и допускает формальное ограничение внутренних контуров;",
        "применение ИИ обосновано для медленной адаптации параметров и ограниченного выбора управляющего действия, а не для бесконтрольной замены быстрого контура;",
        "сравнение должно выполняться со специально настроенными FOC-SVPWM, DTC и FCS-MPC при одинаковых ограничениях и вычислительном бюджете;",
        "научная задача состоит в совместном уменьшении ошибки регулирования, потерь и коммутационной нагрузки при доказуемом сохранении безопасного состояния."], numbered=True)


def add_chapter2(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("2 МАТЕМАТИЧЕСКИЕ МОДЕЛИ И КРИТЕРИИ КАЧЕСТВА", level=1)
    doc.add_heading("2.1 Каталожный профиль и расчётная модель двигателя АИР56В2", level=2)
    add_body(doc, "Для целевого двигателя IEK АИР56В2 в каталожном профиле приняты номинальная мощность 0,25 кВт, напряжение 220/380 В при соединении Δ/Y, ток 1,24/0,72 А, частота 50 Гц и скорость 2720 об/мин. Оператор подтвердил соответствие этих данных выбранному исполнению, однако фотография шильдика конкретного экземпляра и её хэш в пакет доказательств пока не включены. Для стенда с линейным напряжением инвертора 220 В обмотки соединяются треугольником. Внутренняя трёхфазная модель MCSDK использует эквивалентное фазное напряжение 220/√3=127,017 В. Это математическое представление не означает физическое переключение двигателя в звезду.")
    add_equation(doc, "Uph,eq = Uline / √3 = 220 / √3 = 127,017 В", "2.1")
    add_body(doc, "Для одной пары полюсов синхронная скорость при 50 Гц равна 3000 об/мин, а номинальное скольжение составляет (3000-2720)/3000=0,0933. Значение согласуется с двухполюсным исполнением двигателя и используется для проверки ошибок конфигурации.")
    add_equation(doc, "nₛ = 60f₁/p;    s = (nₛ - n)/nₛ", "2.2")
    add_body(doc, "Каталожные и расчётные величины сведены в таблицу 2.1, а положение номинальной точки относительно синхронной скорости показано на рисунке 2.1. Полный машинно-читаемый профиль приведён в приложении А.")
    add_table(doc, "Таблица 2.1 - Каталожные и расчётные параметры АИР56В2", ["Параметр", "Значение", "Статус"], [
        ["Мощность Pn", "250 Вт", "Каталог; фото экземпляра не приложено"],
        ["Напряжение Δ/Y", "220/380 В", "Каталог; подтверждено оператором"],
        ["Ток Δ/Y", "1,24/0,72 А", "Каталог; подтверждено оператором"],
        ["Частота", "50 Гц", "Каталог; подтверждено оператором"],
        ["Скорость", "2720 об/мин", "Каталог; подтверждено оператором"],
        ["Пары полюсов", "1", "Расчёт по 50 Гц и 2720 об/мин"],
        ["Rs, Rr, Lls, Llr, Lm, J", "не измерены", "IOC demo-значения запрещены для FOC"],
    ], [3600, 2200, 3838], font_size=10, page_break_before=True)
    add_figure(doc, figs["speed"], "Рисунок 2.1 - Связь частоты и скорости для двигателя с одной парой полюсов")

    doc.add_heading("2.2 Модель асинхронной машины в координатах dq", level=2)
    add_body(doc, "Для построения FOC используется модель в синхронно вращающейся системе координат dq. Напряжения статора связаны с токами и потокосцеплениями выражениями")
    add_equation(doc, "uₛd = Rₛiₛd + dψₛd/dt - ωₑψₛq", "2.3")
    add_equation(doc, "uₛq = Rₛiₛq + dψₛq/dt + ωₑψₛd", "2.4")
    add_body(doc, "Уравнения ротора для короткозамкнутой машины записываются при нулевых приложенных напряжениях ротора. Электрическая скорость скольжения равна разности синхронной скорости системы координат и электрической скорости ротора.")
    add_equation(doc, "0 = Rᵣiᵣd + dψᵣd/dt - ωslψᵣq", "2.5")
    add_equation(doc, "0 = Rᵣiᵣq + dψᵣq/dt + ωslψᵣd", "2.6")
    add_body(doc, "Связь потокосцеплений и токов определяется собственными и взаимной индуктивностями. Для вычислительной реализации параметры приводятся к одной стороне и проверяются тестами размерности. Механическая часть описывается уравнением баланса моментов.")
    add_equation(doc, "J dωm/dt = Mₑ - Mload - Bωm", "2.7")
    add_equation(doc, "Mₑ = (3/2)p(ψₛd iₛq - ψₛq iₛd)", "2.8")
    add_body(doc, "В тяговой задаче нагрузочный момент не является постоянным: он зависит от сопротивления движению, передаточного отношения, радиуса колеса и сцепления. В лабораторной работе нагрузка задаётся тормозной машиной или моделью HIL. Для переноса результатов к подвижному составу механическая модель дополняется массой поезда, уклоном, кривыми и ограничением силы тяги по сцеплению.")

    doc.add_heading("2.3 Скалярный закон, номинальная область и ослабление поля", level=2)
    add_body(doc, "В основной области частота изменяется от нуля до 50 Гц, а напряжение возрастает от Uboost до номинального. После достижения доступного напряжения дальнейший рост частоты возможен только в режиме ослабления поля. При приблизительно постоянной мощности допустимый момент уменьшается обратно пропорционально скорости. Такой режим не включён в текущую прошивку и рассматривается как последующая исследовательская стадия.")
    add_equation(doc, "U₁(f) = Uboost + (UN - Uboost)f/fN,    0 ≤ f ≤ fN", "2.9")
    add_equation(doc, "U₁(f) = UN;    ψ ~ UN/f;    Mmax ~ 1/f,    f > fN", "2.10")
    add_figure(doc, figs["vf"], "Рисунок 2.2 - Расчётный закон V/f и перспективная область частот выше 50 Гц")
    add_body(doc, "На рисунке 2.2 область до 50 Гц соответствует текущему программному ограничению, а горизонтальный участок напряжения обозначает только будущую область ослабления поля.")
    add_body(doc, "Переход выше 50 Гц допускается только после проверки механической предельной скорости двигателя и нагрузки, напряжения звена, тока, нагрева и устойчивости регуляторов. Простое снятие программного ограничения без этих данных недопустимо. В экспериментальном плане частотный диапазон расширяется ступенчато с независимым контролем вибрации и температуры.")

    doc.add_heading("2.4 Преобразования Clarke и Park, структура FOC", level=2)
    add_body(doc, "Трёхфазные токи преобразуются в неподвижные координаты αβ, а затем поворачиваются на оцененный электрический угол. При симметричной системе достаточно двух измеренных фазных токов, но реконструкция третьего должна учитывать синхронность выборки и насыщение АЦП.")
    add_equation(doc, "iα = ia;    iβ = (ia + 2ib)/√3", "2.11")
    add_equation(doc, "id = iα cos θe + iβ sin θe", "2.12")
    add_equation(doc, "iq = -iα sin θe + iβ cos θe", "2.13")
    add_body(doc, "При ориентации оси d по потокосцеплению ротора ψrq стремится к нулю. Команда isd задаёт поток, а isq — момент. Внешний регулятор скорости формирует isq*, внутренние регуляторы токов формируют напряжения usd* и usq*, после чего обратное преобразование и SVPWM создают коэффициенты заполнения.")
    add_equation(doc, "Mₑ ≈ (3/2)p(Lm/Lr)ψrd isq", "2.14")
    add_body(doc, "Функциональные связи преобразований, регуляторов, наблюдателя и модулятора представлены на рисунке 2.3.")
    add_figure(doc, figs["foc"], "Рисунок 2.3 - Функциональная структура FOC с оцениванием потока")

    doc.add_heading("2.5 Пространственно-векторная модуляция и ограничения", level=2)
    add_body(doc, "Двухуровневый трёхфазный инвертор имеет восемь комбинаций ключей: шесть активных векторов и два нулевых. Классическая SVPWM аппроксимирует заданный вектор в пределах периода переключения соседними активными и нулевыми состояниями. При прямом интеллектуальном выборе конечного набора векторов тот же физический набор становится пространством действий.")
    add_equation(doc, "V*Ts = VkT₁ + Vk+1T₂ + V₀T₀;    T₀ + T₁ + T₂ = Ts", "2.15")
    add_body(doc, "Геометрия допустимых активных векторов и произвольного задания V* показана на рисунке 2.4.")
    add_figure(doc, figs["svpwm"], "Рисунок 2.4 - Допустимые пространственные векторы напряжения инвертора")
    add_body(doc, "Ограничения модуляции включают запрет одновременного включения верхнего и нижнего ключей плеча, минимальный импульс, dead-time, максимальный коэффициент модуляции, токовый предел и тепловой предел. Safety Gateway проверяет действие независимо от алгоритма, сформировавшего его. При невозможности принять действие используется заранее определённый fallback или отключение PWM.")
    add_body(doc, "Для сопоставления модели со стендом идеальный инвертор должен быть дополнен падением напряжения на ключах и диодах, искажением dead-time, задержкой обновления PWM и ограничением напряжения звена. Измерительный тракт описывается коэффициентами передачи, смещением АЦП, квантованием, шумом и задержкой синхронной выборки. Температурная подсистема задаёт дрейф сопротивлений и допустимые пределы нагрева. Эти неидеальности образуют множество неопределённости, в котором должны проверяться базовый FOC и адаптивный метод [21-24, 27-29].")

    doc.add_heading("2.6 Критерии энергетической эффективности и робастности", level=2)
    add_body(doc, "КПД привода вычисляется по синхронным измерениям электрической и механической мощности с явным учётом направления потока энергии. Для малой машины абсолютная погрешность датчиков может быть сопоставима с исследуемой разностью алгоритмов, поэтому необходима предварительная оценка неопределённости.")
    add_equation(doc, "ηmot=Pmech/Pdc;  ηreg=|Pdc|/|Pmech|;  Ecycle=∫Pdc dt", "2.16")
    add_body(doc, "Первое отношение применяется только в двигательном режиме Pdc>0 и Pmech>0, второе — при рекуперативном торможении с согласованным знаком мощности. Для полного тягового цикла основной энергетической метрикой служит интеграл Ecycle и удельная энергия, отнесённая к механической работе, пути либо заданному циклу. Участки с различным направлением мощности не усредняются одной формулой КПД.")
    add_body(doc, "Дополнительные показатели включают среднеквадратический ток, интеграл квадрата ошибки скорости, пиковый ток, пульсации момента, коммутационные события и температуру. Для каждого сценария задаётся одинаковая длительность, начальное состояние и профиль нагрузки. Статистический вывод строится по серии повторов с доверительными интервалами, а не по одной красивой осциллограмме.")
    add_equation(doc, "IAE = ∫|n* - n|dt;    I²t = ∫is²dt", "2.17")

    doc.add_heading("2.7 Граница переноса к тяговому электроприводу", level=2)
    add_body(doc, "Лабораторный АИР56В2 не является тяговым двигателем, поэтому одной нормировки электрических величин недостаточно для отраслевого вывода. Перенос должен выполняться через отдельную модель механической части поезда, включающую приведённую массу, передаточное отношение, радиус колеса, КПД передачи, сопротивление движению, уклон и ограничение силы тяги по сцеплению.")
    add_equation(doc, "Meq·dv/dt = Ft − (A+Bv+Cv²) − Meq·g·i;    Ft = ηg·ig·Mm/rw", "2.18")
    add_equation(doc, "|Ft| ≤ μ(v,weather)·Ma·g;    Especific = ∫Pdc dt / S", "2.19")
    add_body(doc, "Коэффициенты A, B, C, приведённая масса Meq, сцепной вес Ma, коэффициент сцепления μ и профиль пути должны задаваться для выбранной серии подвижного состава. В рабочей редакции численные параметры тягового объекта не выбраны, поэтому модель (2.18)-(2.19) определяет обязательную структуру будущего эксперимента, но не даёт количественного результата. До заполнения этой модели выводы относятся к лабораторному асинхронному электроприводу и методике его последующего масштабирования.")

    doc.add_heading("2.8 Выводы по главе 2", level=2)
    add_body(doc, "По результатам математического описания установлено:")
    add_list(doc, [
        "принятые каталожные данные АИР56В2 задают числовую часть кандидата V/f-профиля для физического соединения обмоток треугольником при 220 В, тогда как 127,017 В является только эквивалентным фазным напряжением модели; профиль должен быть подтверждён по конкретному экземпляру;",
        "параметры Rs, Rr, Lls, Llr, Lm, J и B не следуют из шильдика и должны быть получены экспериментальной идентификацией до настройки FOC;",
        "dq-модель, преобразования Clarke и Park и модель двухуровневого инвертора образуют единый цифровой двойник для сравнения классических и интеллектуальных алгоритмов;",
        "диапазон выше 50 Гц относится к ослаблению поля и требует отдельного механического, теплового и электрического допуска;",
        "критерий качества должен одновременно учитывать КПД, динамическую ошибку, токовую нагрузку, коммутационные события и нарушения ограничений."], numbered=True)


def add_chapter3(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("3 ПРОГРАММНО-АППАРАТНАЯ ПЛАТФОРМА И ПРОСЛЕЖИВАЕМОСТЬ", level=1)
    doc.add_heading("3.1 Архитектура стенда", level=2)
    add_body(doc, "Проект платформы построен по иерархическому принципу. NUCLEO-G431RB должен выполнять быстрый контур MCSDK, управлять таймерами, АЦП и аппаратным входом break. X-NUCLEO-IHM09M2 согласует стандартный motor-control connector с платой STEVAL-IPM15B [19-24]. UNO Q предназначен для функций верхнего уровня, журналирования и будущей интеллектуальной адаптации, но не должен управлять шестью PWM-сигналами непосредственно [25].")
    add_body(doc, "Разделение функций показано на рисунке 3.1, а назначение и критические ограничения компонентов приведены в таблице 3.1.")
    add_figure(doc, figs["architecture"], "Рисунок 3.1 - Функциональная архитектура стенда и разделение ответственности")
    add_table(doc, "Таблица 3.1 - Основные компоненты платформы", ["Компонент", "Назначение", "Критическое ограничение"], [
        ["STM32G431RBTx", "Контур реального времени", "Плата ожидается с RBT6; IOC пока RBT3"],
        ["NUCLEO-G431RB", "Отладочная плата", "Прошивка через встроенный ST-Link"],
        ["X-NUCLEO-IHM09M2", "Переходник", "34-контактный motor-control интерфейс"],
        ["STEVAL-IPM15B", "Трёхфазная IPM-плата", "Мощность сверять с ревизией UM2014"],
        ["UNO Q", "Верхний уровень и ИИ", "Только ограниченный UART"],
        ["ISO7721", "Развязка UART", "Раздельное питание сторон 3,3 В"],
        ["АИР56В2", "Объект управления", "220 В Δ; 1,24 А; 50 Гц [26]"],
    ], [2500, 3300, 3838], font_size=9, page_break_before=True)

    doc.add_heading("3.2 Конфигурация MCSDK и двигателя", level=2)
    add_body(doc, "Основной проект расположен в каталоге mcsdk_reference/.../AIR56B2_..._NOT_FOR_HV. Он создан из официального примера ACIM V/F Open Loop для NUCLEO-G431RB, X-NUCLEO-IHM09M2 и STEVAL-IPM15B. Полный путь приведён в приложении А. Метаданные содержат MC_WORKBENCH_VERSION=6.2.0, WB_to_Mx_version=6.4.2, CUBE_MX_VER=6.16.1, MxCube.Version=6.18.1 и пакет STM32CubeG4 1.6.3; эти поля сохраняются вместе, поскольку отражают разные этапы генерации.")
    add_body(doc, "На ожидаемой физической плате NUCLEO-G431RB должен быть проверен STM32G431RBT6, тогда как импортированный IOC содержит Mcu.CPN=STM32G431RBT3 и общее обозначение STM32G431R(6-8-B)Tx. До осмотра платы и следующей контролируемой регенерации используется семейное обозначение STM32G431RBTx; расхождение метаданных должно быть устранено, а бинарник рассматривается только как кандидат.")
    add_body(doc, "В конфигурации заданы POLE_PAIR_NUM=1, NOMINAL_FREQ=50 Гц, MOTOR_MAX_SPEED_RPM=3000, эквивалентное фазное напряжение 127 В, ток 1,24 А и коэффициент потока FLUX_K=0,5717009 В·с. Сборка Debug завершена успешно с сохранением журнала и манифеста. Конфигурация намеренно содержит в названии NOT_FOR_HV, поскольку не является допуском к высоковольтному пуску.")
    add_body(doc, "Повторная генерация выполняется только в копию проекта. Это защищает пользовательский UART-адаптер от перезаписи Workbench/CubeMX и позволяет сравнить отчёт регенерации. После каждой генерации автоматически проверяются версия протокола, частотное ограничение, запрещённые режимы, отсутствие прямой записи в PWM и согласованность артефактов.")
    add_body(doc, "Аудит показал, что часть силовых и моторных констант унаследована от исходного примера и ещё не подтверждена измерениями. В частности, IOC содержит Rs=2,85 Ом, а сгенерированный acim_motor_parameters.h — RS=0,35 Ом; оба значения запрещено интерпретировать как измеренное сопротивление АИР56В2. Текущие параметры сведены в таблицу 3.2.")
    add_table(doc, "Таблица 3.2 - Фактические параметры текущей конфигурации MCSDK", ["Параметр", "Значение в проекте", "Статус"], [
        ["Частота PWM", "16 кГц", "IOC и drive_parameters.h; стенд не проверен"],
        ["Dead-time", "HW 2000 нс; SW 1350 нс", "Унаследовано; сверить с IPM и осциллограммой"],
        ["Номинальное Udc", "325 В", "Унаследованная расчётная настройка"],
        ["Порог перенапряжения", "390 В", "IOC/drive_parameters.h; аппаратно не проверен"],
        ["Порог пониженного Udc", "10 В", "IOC/drive_parameters.h; проверить назначение"],
        ["Низкочастотная добавка", "20 В до порога 10 Гц", "Не настроена для целевого двигателя"],
        ["Сопротивление статора Rs", "IOC 2,85 Ом; header 0,35 Ом", "Противоречие; измерить до FOC"],
    ], [2800, 2800, 4038], font_size=9)

    doc.add_heading("3.3 Контракт связи UNO Q — Nucleo", level=2)
    add_body(doc, "Проектируемое физическое соединение выполняется через двухканальный цифровой изолятор: UNO D1/Serial1 TX соединяется через канал изолятора с PB7/USART1_RX, а PB6/USART1_TX — с UNO D0/Serial1 RX. Обе стороны питаются от собственных 3,3 В; земли через изолятор не объединяются. Параметры UART: 115200 бод, 8 бит данных, без чётности, один стоповый бит. Каноническая плата hardware/mic_ai_rev2 использует Blue Pill PA2/PA3, поэтому вариант Nucleo по PB6/PB7 пока является схемой миграции и должен быть собран и проверен отдельно. Таблица физических соединений приведена в приложении Б.")
    add_body(doc, "Формат кадра показан на рисунке 3.2, а назначение байтов и проверки приведены в таблице 3.3.")
    add_figure(doc, figs["uart"], "Рисунок 3.2 - Структура 32-байтового командного кадра")
    add_table(doc, "Таблица 3.3 - Поля командного кадра версии 0x02", ["Байты", "Содержание", "Проверка"], [
        ["0-1", "AA 55", "Синхронизация"],
        ["2", "0x02", "Точное совпадение версии"],
        ["3", "0x01 START; 0x08 CLEAR; 0x02 ESTOP", "Взаимоисключающие команды"],
        ["4", "3 SCALAR; 0 OFF", "Иные режимы запрещены"],
        ["5", "sequence", "Возвращается в ответе"],
        ["6-9", "частота, mHz, little-endian", "START: 1-50000 mHz"],
        ["10-13", "историческая амплитуда", "Игнорируется MCSDK"],
        ["14-30", "резерв", "Только нули"],
        ["31", "XOR CRC", "XOR байтов 0-30"],
    ], [1400, 4800, 3438], font_size=9)
    add_body(doc, "При неверной версии, CRC, режиме, ненулевом резерве, частоте выше лимита, E-stop или тишине более 300 мс адаптер вызывает MC_StopMotor1() и защёлкивает отказ. Снятие отказа допускается только чистым кадром CLEAR. Поля токов, температуры и энкодера в текущем ответе равны нулю, поскольку не согласован отдельный формат телеметрии. Выдуманные измерения не передаются.")

    doc.add_heading("3.4 Безопасное состояние и аппаратные защиты", level=2)
    add_body(doc, "Переходы между безопасным, готовым, рабочим и аварийным состояниями представлены на рисунке 3.3.")
    add_figure(doc, figs["states"], "Рисунок 3.3 - Автомат безопасного запуска и остановки")
    add_body(doc, "Безопасное состояние определяется как запрещённая генерация PWM и отсутствие разрешения на силовой пуск. Команда START по пользовательскому протоколу принимается только при валидной связи, допустимой частоте и отсутствии защёлкнутого отказа; известные альтернативные программные пути запуска MCSDK и кнопка PC13 переведены в режим stop-only. Флаг ESTOP в UART является программной командой. Вход break/OCP действует независимо от UART, но отдельная нормально-замкнутая аппаратная цепь E-stop ещё не реализована. При потере связи программный адаптер должен вызвать остановку, а время реакции предстоит измерить на плате.")
    add_body(doc, "Архивная схема hardware/mic_ai_rev2 относится к Blue Pill и содержит неподтверждённый узел реле предзаряда; она помечена как непригодная для сборки текущей конфигурации Nucleo. Для NUCLEO-G431RB — IHM09M2 — IPM15B пока отсутствует выпущенная полная электрическая схема с проверенным 34-контактным соответствием, ERC/DRC, предзарядом, разрядом и аппаратным E-stop. До её выпуска любые таблицы соединений являются проектом миграции, а не монтажным документом.")
    add_body(doc, "Для допуска к силовому запуску требуется реальный выход Nucleo на драйвер предзаряда, контроль достижения Udc, запрет PWM до готовности, размыкание при fault/E-stop/timeout и HIL-проверка. В текущем проекте интерлок не реализован, поэтому release-gate намеренно не разрешает высоковольтный статус.")
    add_body(doc, "Работы со звеном около 310 В DC требуют изолированного источника или штатного ЛАТР с защитой, предохранителей, разрядного резистора, закрытого корпуса, дистанционного аварийного отключения и проверки отсутствия напряжения перед касанием. Программная проверка не снижает электрическую опасность; структура защит и испытаний должна учитывать руководство силовой платы и применимые принципы IEC 61800 [21, 27, 28]. Принципы построения safety case сопоставляются с ISO 26262 только методически [30]; этот стандарт не подменяет железнодорожные нормативы.")

    doc.add_heading("3.5 Сборка, хэширование и управление конфигурацией", level=2)
    add_body(doc, "Команда tools/build_firmware_bundle.ps1 собирает Nucleo и UNO Q, выполняет профильные проверки, формирует манифесты и сверяет SHA-256. Для Nucleo сохраняются ELF, BIN и HEX; для UNO Q дополнительно создаётся upload-payload *.elf-zsk.bin. Манифесты текущей версии подтверждают состав, размер и целостность готовых файлов, но ещё не содержат хэши профиля, исходников, Git commit, компилятора и команды сборки; следовательно, они не являются полной цепочкой происхождения.")
    add_body(doc, "Минимальный набор фиксируемых сведений приведён в таблице 3.4.")
    add_table(doc, "Таблица 3.4 - Правила воспроизводимости", ["Объект", "Обязательная фиксация", "Причина"], [
        ["Двигатель", "Марка, фото шильдика, схема Δ/Y", "Снизить риск чужого профиля"],
        ["MCSDK-проект", "Версии MCSDK/CubeMX/CubeG4", "Повторяемая генерация"],
        ["Прошивка", "Размер и SHA-256", "Однозначность образа"],
        ["Протокол", "Версия и тестовые векторы", "Совместимость узлов"],
        ["Эксперимент", "Сценарий, seed, режим, калибровка", "Повторяемость результата"],
        ["Результат", "Сырые данные и отчёт", "Проверяемость вывода"],
    ], [1900, 4300, 3438], font_size=10)

    doc.add_heading("3.6 Готовность к расширению", level=2)
    add_body(doc, "Переход к FOC выполняется после идентификации параметров и сохранения отдельного профиля. Добавление датчика положения AS5600 возможно как исследовательского канала низкой скорости, однако он не должен автоматически считаться тяговым датчиком. Для оценки бездатчикового алгоритма желательно иметь независимый эталонный энкодер.")
    add_body(doc, "UNO Q может выполнять предобработку данных, обучение малой модели, планирование тестов и адаптацию медленных параметров. Быстрый токовый контур и защита остаются на STM32G431. Такая декомпозиция позволяет добавлять функции без изменения временных гарантий MCSDK.")
    doc.add_heading("3.7 Выводы по главе 3", level=2)
    add_body(doc, "Разработка программно-аппаратной платформы дала следующие результаты:")
    add_list(doc, [
        "сформирован логический проект цепочки NUCLEO-G431RB, X-NUCLEO-IHM09M2, STEVAL-IPM15B и АИР56В2 с разделением быстрого и верхнего уровней; физическая совместимость монтажа ещё не подтверждена;",
        "числовая часть V/f-конфигурации согласована с каталожным профилем и программно ограничена частотой 50 Гц; происхождение данных конкретного экземпляра ещё не приложено;",
        "программный UART-контракт версии 0x02 и схема будущей изоляции запрещают режимы FOC и DUTY, контролируют CRC, резервные поля и тайм-аут 300 мс;",
        "сборочные артефакты идентифицированы размерами и SHA-256; для доказательства происхождения требуется добавить хэши профиля, исходников и toolchain;",
        "аппаратный E-stop, предзаряд, реальные сигналы break/OCP и параметры двигателя для FOC остаются открытыми gate-условиями силового запуска."], numbered=True)
    add_body(doc, "Следовательно, на текущем этапе подготовлены программная конфигурация, контракт взаимодействия и неполный проект соединений, но не выпущенная электрическая схема и не физически проверенная платформа. Первый аппаратный цикл должен подтвердить ревизии плат, ориентацию 34-контактного кабеля, фактическую распиновку, уровни 3,3 В и отсутствие конфликта UART с отладочными функциями.")
    add_body(doc, "До подачи напряжения звена требуется устранить расхождение обозначения корпуса STM32G431 в IOC, реализовать предзаряд и разряд, проверить аппаратный break и OCP, а затем повторно сформировать комплект артефактов с журналом сборки. Эти условия перенесены в gate-критерии главы 5 и чек-лист приложения Г.")


def add_chapter4(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("4 КОНЦЕПЦИЯ МЕТОДА ОГРАНИЧЕННОЙ ИНТЕЛЛЕКТУАЛЬНОЙ АДАПТАЦИИ FOC", level=1)
    doc.add_heading("4.1 Иерархическая организация управления", level=2)
    add_body(doc, "Предлагаемый метод разделяет вычисления по критичности и характерному времени. В предлагаемой архитектуре нижний уровень должен реализовать FOC, SVPWM, синхронную выборку токов, break и watchdog на STM32G431. Средний уровень должен оценивать медленно изменяющиеся параметры и корректировать ограниченный набор коэффициентов. Верхний уровень UNO Q предназначен для обучения моделей, анализа серий опытов и формирования предложений, которые не становятся действием без проверки.")
    add_body(doc, "Иерархия вычислительных уровней и направленность ограниченных команд представлены на рисунке 4.1.")
    add_figure(doc, figs["hierarchy"], "Рисунок 4.1 - Три уровня предлагаемой системы управления")
    add_body(doc, "Такая архитектура отличается от прямого нейросетевого PWM тем, что обучаемая модель не является единственным источником устойчивости. При низкой уверенности, выходе за область обучения, превышении времени расчёта или конфликте ограничений используется классический FOC-SVPWM. Переход на fallback фиксируется как отдельная метрика.")

    doc.add_heading("4.2 Адаптация параметров модели и регуляторов", level=2)
    add_body(doc, "Сопротивления и индуктивности влияют на оценку потока и развязку токовых каналов. Вектор параметров θ=[Rs, Rr, Lm, Lσ, J, B] оценивается по измеренным напряжениям, токам и скорости в разрешённых режимах. Допустимое множество Θ задаётся покомпонентными интервалами, полученными из идентификации и оценки неопределённости, а не произвольными процентами. Обновление не должно быть быстрее физически возможного изменения параметра.")
    add_equation(doc, "γk = 1{rk>rmin ∧ μk>μmin ∧ ck>cmin ∧ Δtk>tdwell};   θk+1 = ProjΘ{θk + γkKkrk}", "4.1")
    add_body(doc, "Событие обновления γk разрешается только при нормированном остатке rk выше порога чувствительности, достаточной информативности μk регрессионной матрицы, доверии ck и выдержанном интервале tdwell. Матрица Kk должна вычисляться ограниченным RLS- или градиентным законом с нормированием и максимальным шагом; конкретный вариант и численные пороги определяются после идентификации стенда. При нарушении любого условия γk=0 и оценка замораживается.")
    add_body(doc, "Оператор Proj_Theta обеспечивает инвариант θk∈Θ при θ0∈Θ независимо от необработанного предложения адаптера. Этот простой инвариант не является доказательством устойчивости всей замкнутой системы: до защиты требуется отдельно установить условия ограниченности сигналов и устойчивости FOC с переключаемой оценкой. Коэффициенты PI пересчитываются по проверенным формулам и проходят ограничения по полосе и максимальному управляющему воздействию.")
    add_equation(doc, "Kp,i = Lσ ωci;    Ki,i = Req ωci", "4.2")
    add_body(doc, "Интеллектуальная модель может формировать поправку к θ или к коэффициентам, но базовая физическая модель сохраняется. Это уменьшает объём требуемых данных и обеспечивает интерпретируемость. Обучение проводится с domain randomization в диапазонах неопределённости, а валидация — на отдельных траекториях и параметрах.")

    doc.add_heading("4.3 Перспективное расширение: горизонтный выбор векторов", level=2)
    add_body(doc, "Горизонтный выбор векторов не входит в основной предмет доказывания настоящей работы и рассматривается как перспективное расширение после подтверждения ограниченной адаптации FOC. На каждом шаге формируется конечное множество допустимых последовательностей векторов на горизонте H. Цифровой двойник прогнозирует ток, поток, момент и тепловой прокси. Нейросетевая часть корректирует функцию стоимости или остаточную ошибку модели. Выбирается действие с минимальной стоимостью, прошедшее Safety Gateway.")
    add_equation(doc, "uₖ* = arg min {Σ J(x̂ₖ₊ₕ, uₖ₊ₕ₋₁)},    u ∈ Usafe", "4.3")
    add_body(doc, "Горизонт H=1 близок к одношаговому FCS-MPC. Увеличение горизонта позволяет учитывать будущие коммутации и ограничения, но экспоненциально увеличивает поиск. Для STM32 возможны предварительное сокращение множества кандидатов, beam search, таблицы переходов или вычисление предложения на UNO Q при сохранении локальной валидации Nucleo.")
    add_body(doc, "Коммутационный штраф должен учитывать не только число переключений, но и состояние плеча, ток в момент переключения и тепловой баланс. При одинаковой ошибке тока предпочтительно действие с меньшими ожидаемыми потерями, если оно не ухудшает устойчивость.")

    doc.add_heading("4.4 Safety Gateway и проверяемые инварианты", level=2)
    add_body(doc, "Safety Gateway проектируется как логически обособленный детерминированный модуль с неизменяемой во время опыта логикой проверок и параметрами конкретной силовой платы. Его программная обособленность не означает независимость по общему питанию, тактированию, памяти или отказам микроконтроллера. Входом служат предложенный вектор, текущее состояние, измерения и таймеры; выходом — принятое действие, fallback или отключение. Таблица решений и fault-injection являются тестовым покрытием, но не заменяют формального доказательства модели состояний.")
    add_body(doc, "Перечень инвариантов, условий контроля и реакций приведён в таблице 4.1.")
    add_table(doc, "Таблица 4.1 - Основные инварианты Safety Gateway", ["Инвариант", "Условие", "Реакция"], [
        ["Нет сквозного тока", "Верхний и нижний ключ плеча не включены вместе", "Отклонить; latch fault"],
        ["Dead-time", "Переход выдерживает минимальную паузу", "Отложить или fallback"],
        ["Ток", "|iphase| ≤ Imax", "Отключить PWM; OCP"],
        ["Напряжение", "Udc в допустимом диапазоне", "Запрет пуска/останов"],
        ["Температура", "Tj,est ≤ Tmax", "Деградация/останов"],
        ["Время", "Расчёт завершён до deadline", "Классический SVPWM"],
        ["Доверие", "x внутри области валидации", "Заморозить ИИ-адаптацию"],
    ], [2600, 4200, 2838], font_size=9)
    add_body(doc, "Инварианты проверяются до применения действия. Даже если алгоритм выдаёт недопустимый необработанный запрос, шлюз обязан предотвратить его передачу в PWM. Отдельно тестируется способность детектора увидеть искусственно введённый запрос сквозного тока и нарушение dead-time, иначе отсутствие инцидентов может означать лишь отсутствие действенного теста.")

    doc.add_heading("4.5 Цифровой двойник и цикл обучения", level=2)
    add_body(doc, "Полный цикл от идентификации до HIL- и стендовой верификации показан на рисунке 4.2.")
    add_figure(doc, figs["ai_pipeline"], "Рисунок 4.2 - Цикл данных, обучения и независимой верификации")
    add_body(doc, "Цифровой двойник состоит из физической dq-модели и ограниченной остаточной коррекции. В архивном host-эксперименте контекст включал состояние, выбранный вектор, напряжение αβ, оценку нагрузки, токовый прокси и параметры θ. Обучающая выборка формировалась при случайном изменении Rs и Rr на ±50 %, Lm на ±20 %, J и B на ±100 %.")
    add_body(doc, "Валидационная метрика остаточного слоя показала небольшое улучшение многошагового прогноза, но ухудшение одношагового RMSE. Поэтому остаточный слой не объявляется готовым; более надёжным результатом остаётся θ-conditioned физический двойник. Этот отрицательный результат важен: он показывает необходимость отбора модели по нескольким горизонтам, а не по одной метрике.")
    add_body(doc, "Количественные результаты архивной host-валидации сведены в таблицу 4.2.")
    add_table(doc, "Таблица 4.2 - Результаты архивной host-валидации цифрового двойника", ["Горизонт", "RMSE модели", "RMSE с residual", "Снижение RMSE"], [
        ["1 шаг", "0,001777", "0,001705", "+4,02 %"],
        ["5 шагов", "0,007962", "0,007546", "+5,23 %"],
        ["10 шагов", "0,014152", "0,013466", "+4,85 %"],
        ["50 шагов", "0,043399", "0,042144", "+2,89 %"],
        ["Отдельный one-step набор", "0,003534", "0,004073", "-15,26 %"],
    ], [1800, 2500, 2500, 2838], font_size=10, page_break_before=True)
    add_body(doc, "Положительное значение в последнем столбце означает уменьшение RMSE, отрицательное — ухудшение. Таблица отражает только один архивный прогон компьютерного моделирования и не является результатом идентификации физического АИР56В2. В архиве не зафиксированы достаточные сведения о seeds, числе повторов и доверительных интервалах. После появления стенда данные должны быть собраны заново с версией набора данных, калиброванными датчиками и раздельными обучающей, валидационной и контрольной последовательностями.")

    doc.add_heading("4.6 Программа сравнения алгоритмов", level=2)
    add_body(doc, "Основное сравнение проводится внутри семейства FOC: фиксированный профиль; профиль после идентификации; непрерывная ограниченная адаптация; событийная адаптация; событийная адаптация с интеллектуальной коррекцией. Для каждого варианта фиксируются одинаковые токовые, напряженческие и тепловые пределы. Горизонтный PWM, DTC и FCS-MPC образуют отдельную перспективную программу и не смешиваются с основной абляцией метода адаптации.")
    add_body(doc, "Абляционные варианты и проверяемые эффекты приведены в таблице 4.3.")
    add_table(doc, "Таблица 4.3 - План абляционных исследований", ["Вариант", "Отключаемый механизм", "Проверяемый эффект"], [
        ["A0", "FOC с исходным фиксированным профилем", "Базовая линия"],
        ["A1", "FOC с офлайн-идентификацией", "Влияние идентификации"],
        ["A2", "Непрерывная ограниченная адаптация", "Эффект адаптации и шум"],
        ["A3", "Событийная адаптация без ИИ", "Эффект события γk"],
        ["A4", "Событийная адаптация с ИИ-коррекцией", "Добавочный эффект ИИ"],
        ["A5", "Принудительная заморозка/fallback", "Безопасность деградации"],
    ], [1400, 4300, 3938], font_size=10, page_break_before=True)
    doc.add_heading("4.7 Выводы по главе 4", level=2)
    add_body(doc, "Предложенный метод характеризуется следующими положениями:")
    add_list(doc, [
        "быстрый контур FOC-SVPWM предполагается реализовать на STM32G431; аппаратный break и watchdog должны оставаться детерминированными и независимыми от ИИ;",
        "обучаемый уровень изменяет только ограниченный вектор параметров или ранжирует конечное множество допустимых действий;",
        "каждое предложение проходит логически обособленный Safety Gateway, проверяющий топологию ключей, dead-time, ток, напряжение, температуру, deadline и область доверия;",
        "при отказе модели, выходе за область обучения или превышении времени расчёта система переходит к заранее проверенному FOC-SVPWM;",
        "архивные host-материалы иллюстрируют реализуемость алгоритмической гипотезы и программу абляций в зафиксированном прогоне, но не подтверждают независимую воспроизводимость, аппаратную готовность и превосходство на физическом объекте."], numbered=True)


def add_chapter5(doc: Document, figs: dict[str, Path]) -> None:
    reports = collect_programmatic_evidence()
    nucleo_manifest = read_json(MCSDK_MANIFEST)
    uno_manifest = read_json(UNO_MANIFEST)
    nucleo_sizes = artifact_sizes(nucleo_manifest)
    uno_sizes = artifact_sizes(uno_manifest)

    def passed(report: dict) -> bool:
        return report.get("pass") is True and report.get("runner_exit_code", 0) == 0

    profile_result = "PASS: числовая согласованность" if passed(reports["profile"]) else "FAIL/нет отчёта"
    contract_result = "PASS: статический контракт исходников" if passed(reports["contract"]) else "FAIL/нет отчёта"
    bundle_result = "PASS: состав и SHA-256" if passed(reports["bundle"]) else "FAIL/нет отчёта"
    release_result = "FAIL: HV запрещён" if reports["release"].get("pass") is False else "НЕОЖИДАННЫЙ СТАТУС"
    build_errors = nucleo_manifest.get("build_errors")
    build_warnings = nucleo_manifest.get("build_warnings")
    if nucleo_manifest.get("build_exit_code") == 0 and build_errors == 0 and build_warnings == 0:
        build_result = "PASS: 0 ошибок, 0 предупреждений; log сохранён"
    elif nucleo_manifest.get("build_exit_code") == 0:
        build_result = f"Сборка успешна; errors={build_errors}, warnings={build_warnings}"
    else:
        build_result = "НЕ ПОДТВЕРЖДЕНО"

    doc.add_heading("5 ПРОГРАММНАЯ ПРОВЕРКА И МЕТОДИКА ЭКСПЕРИМЕНТАЛЬНОГО ИССЛЕДОВАНИЯ", level=1)
    doc.add_heading("5.1 Подтверждённые результаты программной проверки", level=2)
    add_body(doc, "При каждой генерации рабочей редакции три проверки целевой конфигурации AIR56B2 запускаются заново, а их JSON-отчёты сохраняются в docs/dissertation_2_9_3/evidence. Проверка профиля подтверждает только числовую согласованность JSON, IOC, исходников и манифеста, а не происхождение шильдика конкретного двигателя. Проверка контракта является статическим анализом исходников и тестовыми векторами; она дополнительно контролирует блокировку команд MCP START/START_STOP и stop-only обработчик PC13. Отдельный высоковольтный release-preflight ожидаемо остаётся FAIL из-за незакрытых аппаратных условий.")
    add_body(doc, "Результаты отдельных проверок и намеренно незакрытый HV-gate приведены в таблице 5.1.")
    add_table(doc, "Таблица 5.1 - Результаты программных проверок", ["Проверка", "Контролируемые свойства", "Результат"], [
        ["Профиль AIR56B2", "1 пара полюсов, 50 Гц, числовые константы, хэши", profile_result],
        ["Контракт UNO Q - Nucleo", "Кадр 0x02, CRC, timeout, запрет режимов и обходов START", contract_result],
        ["Комплект ELF/BIN/HEX", "Состав артефактов, манифесты, согласованность", bundle_result],
        ["Сборка Nucleo", "Компиляция проекта и сохранённый журнал", build_result],
        ["HV release-preflight", "Предзаряд, происхождение профиля, HV interlock", release_result],
    ], [2500, 4800, 2338], font_size=9)
    add_body(doc, "Программная проверка подтверждает внутреннюю согласованность исходников и артефактов, но не измеряет ток двигателя, фронты PWM, dead-time, напряжение звена или корректность фазировки. По этой причине она не может заменить HIL и стенд.")
    add_body(doc, f"По текущим манифестам комплект Nucleo включает ELF размером {nucleo_sizes.get('ACIM-NUCLEOG431RB-IPM15B-VF_OL.elf', 'не подтверждено')} байт, BIN размером {nucleo_sizes.get('ACIM-NUCLEOG431RB-IPM15B-VF_OL.bin', 'не подтверждено')} байт и HEX размером {nucleo_sizes.get('ACIM-NUCLEOG431RB-IPM15B-VF_OL.hex', 'не подтверждено')} байт. Комплект UNO Q включает ELF размером {uno_sizes.get('UNOQ_MOTOR.ino.elf', 'не подтверждено')} байт, BIN размером {uno_sizes.get('UNOQ_MOTOR.ino.bin', 'не подтверждено')} байт, HEX размером {uno_sizes.get('UNOQ_MOTOR.ino.hex', 'не подтверждено')} байт и загрузочный payload elf-zsk.bin размером {uno_sizes.get('UNOQ_MOTOR.ino.elf-zsk.bin', 'не подтверждено')} байт. Полные SHA-256 сохранены в манифестах и проверяются автоматически.")
    add_body(doc, "Граница между полученными и будущими доказательствами показана на рисунке 5.1.")
    add_figure(doc, figs["evidence"], "Рисунок 5.1 - Разделение подтверждённых и неподтверждённых результатов")

    doc.add_heading("5.2 Результаты архивной host-матрицы", level=2)
    add_body(doc, "Архивный исследовательский прототип содержит 31 сценарий: пуск, изменение нагрузки, реверс, торможение, рекуперация, низкая скорость, ослабление поля, ошибки параметров, нагрев, провалы Udc, шумы, задержки, отказы датчиков и fault injection. В отслеживаемом host-прогоне Safe Neural Horizon PWM H2 не зафиксированы нарушения Safety Gateway или неожиданные защёлкивания отказа. Это проверка охвата сценариев, а не достаточное статистическое доказательство: для исследовательского вывода требуется фиксировать seeds, число повторов, разделение train/validation/test и доверительные интервалы.")
    add_body(doc, "Повторная проверка текущего дерева исходников дала 21 успешно пройденный тест для модулей циклической робастной жизнеспособности и конформного охвата. Полная архивная suite Safe Neural Horizon PWM не собирается из-за отсутствия двенадцати вспомогательных модулей формирования отчёта, упаковки и аудита; suite идентификации также зависит от отсутствующих legacy-моделей induction_motor и inverter_ideal. Поэтому численные результаты 31 сценария рассматриваются как исторические файлы свидетельств, а не как независимо воспроизведённый в настоящей редакции расчёт. До научной публикации недостающие исходники должны быть восстановлены либо все серии должны быть повторно сформированы новым воспроизводимым конвейером.")
    add_body(doc, "Проверка специально включала необработанный запрос сквозного тока и эмуляцию перехода без dead-time. Детекторы сработали, а шлюз не передал опасное действие в разрешённый PWM. Этот результат подтверждает логику программного шлюза на ПК, но не временные характеристики MCU и не работу аппаратного break.")
    add_body(doc, "В архиве присутствуют сильные базовые алгоритмы и ограниченные подборы параметров, однако авторский аудит запрещает утверждать универсальное превосходство, готовность к MCU/HIL/стенду или оптимальность обученной модели. В диссертации эти ограничения сохранены как часть методологии научной добросовестности.")
    add_body(doc, "Допустимая интерпретация каждого уровня доказательств приведена в таблице 5.2.")
    add_table(doc, "Таблица 5.2 - Интерпретация уровней доказательств", ["Уровень", "Что подтверждает", "Чего не подтверждает"], [
        ["Static/build", "Синтаксис, контракт, состав", "Физику и тайминги"],
        ["Host simulation", "Логику и сценарии", "Аппаратный привод"],
        ["HIL", "Реальное время и интерфейсы", "Мощность и тепловой режим"],
        ["Низковольтный стенд", "Фазировку и защиты", "Номинальную нагрузку"],
        ["Силовой стенд", "Энергетику и динамику", "Тяговый масштаб без пересчёта"],
        ["Тяговая установка", "Отраслевую применимость", "Иные серии привода"],
    ], [1900, 3750, 3988], font_size=10, page_break_before=True)

    doc.add_heading("5.3 Последовательность HIL- и стендовых испытаний", level=2)
    add_body(doc, "Последовательность этапов и запрет перехода через незакрытый gate представлены на рисунке 5.2.")
    add_figure(doc, figs["test_steps"], "Рисунок 5.2 - Этапы экспериментальной верификации с критериями перехода")
    add_body(doc, "Этап S0 выполнен только в программной и статической части. До активного низковольтного PWM допускаются лишь пассивная проверка монтажа, прозвонка без питания и сверка ревизий плат. На S1 платы соединяются без силового J7; после утверждения полной схемы осциллографом и логическим анализатором проверяются UART, PWM, последовательность запуска и реакция на timeout. На S2 HIL-имитатор формирует токовые, температурные и аварийные сигналы. Переход к S3 допускается только после проверки отдельного аппаратного E-stop, break, предзаряда и разряда звена.")
    add_body(doc, "Первый вращательный пуск выполняется от регулируемого и ограниченного источника, без механической нагрузки, при минимальной частоте и возможности дистанционного отключения. Проверяются порядок фаз, направление, ток холостого хода, вибрация и нагрев. Затем V/f испытывается по ступеням до 50 Гц. FOC вводится только после идентификации и сверки знаков каналов тока и скорости.")
    add_body(doc, "Интеллектуальный режим включается последним. Сначала он работает в shadow mode: формирует предложение, которое журналируется, но не применяется. После статистической проверки и анализа OOD-состояний разрешаются малые поправки к параметрам. Прямой горизонтный PWM допускается только при работоспособном Safety Gateway и мгновенном fallback на FOC-SVPWM.")

    doc.add_heading("5.4 План экспериментов и измерительные каналы", level=2)
    add_body(doc, "Основные сценарии, воздействия, метрики и число повторов приведены в таблице 5.3; форма первичного протокола дана в приложении В.")
    add_table(doc, "Таблица 5.3 - Основные экспериментальные сценарии", ["Сценарий", "Воздействие", "Основные метрики", "Объём серии"], [
        ["Пуск без нагрузки", "0 → 10 → 50 Гц", "Imax, время, вибрация", "по расчёту N"],
        ["Ступень нагрузки", "0,2 → 0,8 Mn", "IAE, ΔM, I²t", "по расчёту N"],
        ["Сброс нагрузки", "0,8 → 0,2 Mn", "перерегулирование", "по расчёту N"],
        ["Тяговый профиль", "разгон, выбег, торможение, рекуперация", "энергия, сцепление, нагрев", "по расчёту N"],
        ["Изменение Udc", "−10 %, −20 %", "устойчивость, fallback", "по расчёту N"],
        ["Нагрев", "длительный номинал", "Rs, T, КПД", "≥5 циклов"],
        ["Потеря UART", "timeout > 300 мс", "время останова", "≥20 инъекций"],
        ["Отказ датчика", "обрыв/заморозка", "fault, безопасное состояние", "≥20 инъекций"],
        ["Ослабление поля", "50 → 60…80 Гц", "ток, момент, вибрация", "после допуска"],
    ], [2200, 2600, 3100, 1738], font_size=9, page_break_before=True)
    add_body(doc, "Минимальный измерительный набор включает Udc, Idc, два фазных тока, скорость или положение, температуру корпуса двигателя и радиатора, состояния PWM/break, UART и E-stop. Для энергетического сравнения требуется синхронная выборка Udc и Idc, а для механической мощности — датчик момента или калиброванная нагрузочная машина.")
    add_body(doc, "Перед серией фиксируются калибровки, температура окружающей среды, состояние двигателя, версия прошивки и хэш конфигурации. Порядок алгоритмов рандомизируется, чтобы уменьшить влияние нагрева. Если температура вышла за коридор, серия прекращается или данные маркируются как несопоставимые.")

    doc.add_heading("5.5 Статистическая обработка", level=2)
    add_body(doc, "До начала серии фиксируются первичная метрика — изменение КПД Δη, нулевая и альтернативная гипотезы, минимально практически значимый эффект, уровень значимости α=0,05 и требуемая мощность не ниже 0,8. Число повторов рассчитывается по дисперсии пилотной серии, а не выбирается только из удобства эксперимента. Для вторичных метрик заранее задаётся способ контроля множественных сравнений. Баланс неопределённости измерений мощности должен быть меньше заявляемого эффекта.")
    add_body(doc, "Для каждой метрики рассчитываются медиана, среднее, стандартное отклонение, 95%-ный доверительный интервал и худший наблюдённый случай. При парном A/B-дизайне анализируется разность результатов в одинаковом сценарии. Наряду со статистической значимостью указывается практический размер эффекта. При ненормальности парных разностей применяется bootstrap-интервал или заранее выбранный непараметрический критерий.")
    add_equation(doc, "Δη = ηAI − ηFOC;    CI₉₅ = mean(Δη) ± t₀,₉₇₅;N−1 · sΔ/√N", "5.1")
    add_body(doc, "Нарушение ограничения тока, сквозной ток, отказ аппаратной остановки или потеря контроля над PWM рассматриваются как безусловный провал независимо от средней энергетической метрики. Алгоритм с лучшим КПД, но худшей безопасностью не принимается.")
    add_body(doc, "Для спектральных показателей длина окна, частота дискретизации и способ выбора основной гармоники фиксируются заранее. Архивная метрика THD-like по доминирующему FFT-бину не приравнивается к показанию сертифицированного анализатора качества электроэнергии.")

    doc.add_heading("5.6 Критерии завершения этапов", level=2)
    add_body(doc, "Gate-критерии сведены в таблицу 5.4. Чек-лист силового допуска приведён в приложении Г, а связь задач с доказательствами — в приложении Д.")
    add_table(doc, "Таблица 5.4 - Gate-критерии дальнейших работ", ["Переход", "Обязательные условия"], [
        ["S0 → S1", "Сборка PASS; профиль и хэши; выпущена полная схема; пассивная прозвонка"],
        ["S1 → S2", "PWM/timeout и break наблюдаемы; аппаратный E-stop реализован; нет конфликтов пинов"],
        ["S2 → S3", "Предзаряд, break, OCP и discharge прошли fault injection"],
        ["V/f → FOC", "Измерены Rs, Rr, Lls, Llr, Lm, J; датчики откалиброваны"],
        ["FOC → ИИ shadow", "FOC устойчив во всей тестовой области"],
        ["Shadow → ИИ active", "OOD/fallback/Safety Gateway подтверждены HIL"],
        ["50 Гц → field weakening", "Механический, тепловой и электрический допуск"],
    ], [2800, 6838], font_size=10)
    doc.add_heading("5.7 Выводы по главе 5", level=2)
    add_body(doc, "Проведённая программная проверка и разработанная методика эксперимента позволяют сделать следующие выводы:")
    add_list(doc, [
        "числовой профиль АИР56В2, статический контракт UNO Q — Nucleo и комплект прошивок прошли раздельные автоматизированные проверки; отдельный HV release-preflight остаётся в состоянии FAIL;",
        "сборка Nucleo завершена успешно, а журнал, размеры и хэши артефактов зафиксированы в манифесте;",
        "программный PASS подтверждает целостность конфигурации, но не фазировку, dead-time, точность АЦП, предзаряд и электрическую безопасность;",
        "эксперимент разбит на gate-этапы от статической проверки и низковольтного bench до HIL, V/f, FOC и ИИ A/B;",
        "любое нарушение защит считается провалом независимо от средней энергетической эффективности;",
        "архивное host-моделирование используется как подготовительный уровень доказательств и не подменяет физический эксперимент."], numbered=True)


def add_conclusion(doc: Document) -> None:
    doc.add_heading("ЗАКЛЮЧЕНИЕ", level=1)
    add_body(doc, "В рабочей редакции сформулирована научная задача повышения энергетической эффективности и робастности асинхронного тягового электропривода посредством ограниченной адаптации детерминированного FOC. Горизонтное интеллектуальное формирование ШИМ оставлено перспективным расширением и не включено в основной предмет доказывания. Тема соответствует направлениям паспорта специальности 2.9.3, связанным с тяговым приводом, силовыми преобразователями, автоматизацией, энергетическими потерями, моделированием и испытаниями.")
    add_body(doc, "Основные результаты, полученные на текущем этапе, состоят в следующем:")
    add_list(doc, [
        "подготовлены прослеживаемая программная конфигурация и логический проект подключения NUCLEO-G431RB — X-NUCLEO-IHM09M2 — STEVAL-IPM15B — АИР56В2 с верхним уровнем UNO Q; полная выпущенная электрическая схема и физическая платформа отсутствуют;",
        "числовой V/f-профиль согласован с каталожными данными 220 В Δ, 1,24 А, 50 Гц, 2720 об/мин и одной парой полюсов; фото шильдика конкретного экземпляра не приложено; внутри модели MCSDK используется эквивалентное фазное напряжение 127,017 В;",
        "программно реализован 32-байтовый UART-контракт версии 0x02 и разработана схема изолированного подключения; физический тракт ещё не собран и не проверен;",
        "подтверждены успешная сборка Nucleo с журналом 0 ошибок и 0 предупреждений, числовая согласованность проверяемой части V/f-профиля, статическая совместимость контрактов и целостность ELF/BIN/HEX по SHA-256;",
        "сформирована единая математическая основа V/f, dq-модели, FOC и SVPWM, позволяющая использовать одинаковые состояния и критерии в классическом и интеллектуальном управлении;",
        "предложена иерархическая архитектура интеллектуальной адаптации, в которой обучаемая модель не имеет прямого доступа к силовым ключам, а каждая коррекция проверяется детерминированным Safety Gateway с параметрами конкретной платы;",
        "сформирован план сравнительных и абляционных исследований с gate-критериями, одинаковыми ограничениями, повторными опытами и доверительными интервалами;",
        "установлены границы доказанного результата: текущий пакет готов к программным проверкам, осмотру и пассивной прозвонке будущего монтажа; активная низковольтная проверка PWM допускается только после выпуска полной схемы и проверки UART, break/OCP и аппаратного E-stop, а высоковольтный пуск, FOC и активное ИИ-управление запрещены."], numbered=True)
    add_body(doc, "Практическое продолжение работы включает реализацию и HIL-проверку предзаряда, калибровку измерительных каналов, определение Rs, Rr, Lls, Llr, Lm и J, ступенчатые испытания V/f, настройку FOC и последовательное включение ИИ сначала в shadow mode, затем в ограниченном активном режиме. Выход выше 50 Гц рассматривается отдельным экспериментом ослабления поля после механического и теплового допуска.")
    add_body(doc, "Окончательные формулировки научной новизны, практической значимости и положений, выносимых на защиту, должны опираться на эти будущие статистически подтверждённые результаты. До завершения стендовой программы рабочая редакция сохраняет различие между расчётным, программно проверенным и экспериментально доказанным результатом.")


def add_illustrative_material(doc: Document) -> None:
    doc.add_heading("СПИСОК ИЛЛЮСТРАТИВНОГО МАТЕРИАЛА", level=1)

    def entry(text: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=12)

    doc.add_heading("Рисунки", level=2)
    figures = [
        "Рисунок 1 — Логика поэтапного развития исследуемой системы управления",
        "Рисунок 2.1 — Связь частоты и скорости для двигателя с одной парой полюсов",
        "Рисунок 2.2 — Расчётный закон V/f и перспективная область частот выше 50 Гц",
        "Рисунок 2.3 — Функциональная структура FOC с оцениванием потока",
        "Рисунок 2.4 — Допустимые пространственные векторы напряжения инвертора",
        "Рисунок 3.1 — Функциональная архитектура стенда и разделение ответственности",
        "Рисунок 3.2 — Структура 32-байтового командного кадра",
        "Рисунок 3.3 — Автомат безопасного запуска и остановки",
        "Рисунок 4.1 — Три уровня предлагаемой системы управления",
        "Рисунок 4.2 — Цикл данных, обучения и независимой верификации",
        "Рисунок 5.1 — Разделение подтверждённых и неподтверждённых результатов",
        "Рисунок 5.2 — Этапы экспериментальной верификации с критериями перехода",
        "Рисунок Б.1 — Формат командного кадра интерфейса",
    ]
    for item in figures:
        entry(item)

    doc.add_heading("Таблицы", level=2)
    tables = [
        "Таблица 1 — Статус основных частей исследования",
        "Таблица 2 — Принятые сокращения",
        "Таблица 1.1 — Группы требований к системе управления",
        "Таблица 1.2 — Критический обзор методов управления и адаптации",
        "Таблица 2.1 — Каталожные и расчётные параметры АИР56В2",
        "Таблица 3.1 — Основные компоненты платформы",
        "Таблица 3.2 — Фактические параметры текущей конфигурации MCSDK",
        "Таблица 3.3 — Поля командного кадра версии 0x02",
        "Таблица 3.4 — Правила воспроизводимости",
        "Таблица 4.1 — Основные инварианты Safety Gateway",
        "Таблица 4.2 — Результаты архивной host-валидации цифрового двойника",
        "Таблица 4.3 — План абляционных исследований",
        "Таблица 5.1 — Результаты программных проверок",
        "Таблица 5.2 — Интерпретация уровней доказательств",
        "Таблица 5.3 — Основные экспериментальные сценарии",
        "Таблица 5.4 — Gate-критерии дальнейших работ",
        "Таблица А.1 — Содержимое зафиксированного расчётного профиля",
        "Таблица Б.1 — Соединения через ISO7721",
        "Таблица В.1 — Идентификация опыта",
        "Таблица В.2 — Фиксация результатов",
        "Таблица Г.1 — Контроль перед подачей высокого напряжения",
        "Таблица Д.1 — Связь задач, программных объектов и подтверждений",
    ]
    for item in tables:
        entry(item)


def add_references(doc: Document) -> None:
    doc.add_heading("СПИСОК ЛИТЕРАТУРЫ", level=1)
    refs = [
        "Паспорт научной специальности 2.9.3 «Подвижной состав железных дорог, тяга поездов и электрификация» [Электронный ресурс]. URL: https://vak.gisnauka.ru/s3-files/01cc80c69fae4988a0246a8f5e2774e7:fisgna/public/media/uploaded/news_files/4dfe14e2-84dc-45c3-9909-718a368c5fe6/9f3cb2d7-2ce0-4598-845f-ee3298c_EhKPLK2.pdf (дата обращения: 25.08.2026).",
        "ГОСТ Р 7.0.11-2011. Система стандартов по информации, библиотечному и издательскому делу. Диссертация и автореферат диссертации. Структура и правила оформления. М.: Стандартинформ, 2012.",
        "Постановление Правительства Российской Федерации от 24.09.2013 № 842 «О порядке присуждения учёных степеней» (с изм. и доп.).",
        "Регламент о порядке принятия диссертационных работ советами по защите диссертаций ФГБОУ ВО ПГУПС [Электронный ресурс]. URL: https://www.pgups.ru/university/administration/the-scientific-council/dissertation-councils/D_44.2.004.04/ (дата обращения: 25.08.2026).",
        "Сычугов А. Н. Совершенствование систем управления электрическим подвижным составом: дис. ... канд. техн. наук. Санкт-Петербург, 2024.",
        "Ключев В. И. Теория электропривода. М.: Энергоатомиздат, 2001.",
        "Bose B. K. Modern Power Electronics and AC Drives. Upper Saddle River: Prentice Hall, 2002.",
        "Krause P. C., Wasynczuk O., Sudhoff S. D., Pekarek S. Analysis of Electric Machinery and Drive Systems. 3rd ed. Wiley-IEEE Press, 2013.",
        "Blaschke F. The principle of field orientation as applied to the new transvector closed-loop control system for rotating-field machines // Siemens Review. 1972. Vol. 34. P. 217-220.",
        "Holtz J. Sensorless control of induction motor drives // Proceedings of the IEEE. 2002. Vol. 90, no. 8. P. 1359-1394. DOI: 10.1109/JPROC.2002.800726.",
        "Vas P. Sensorless Vector and Direct Torque Control. Oxford: Oxford University Press, 1998.",
        "Kazmierkowski M. P., Krishnan R., Blaabjerg F. Control in Power Electronics: Selected Problems. Academic Press, 2002.",
        "Yang H.-T., Huang K.-Y., Huang C.-L. An artificial neural network based identification and control approach for the field-oriented induction motor // Electric Power Systems Research. 1994. Vol. 30, no. 1. P. 35-45. DOI: 10.1016/0378-7796(94)90057-4.",
        "Sahu A., Mohanty K. B., Mishra R. N. Development and experimental realization of an adaptive neural-based discrete model predictive direct torque and flux controller for induction motor drive // Applied Soft Computing. 2021. Vol. 108. Art. 107418. DOI: 10.1016/j.asoc.2021.107418.",
        "Struharňanský Ľ., Vittek J., Makyš P., Ilončiak J. Vector Control Techniques for Traction Drive with Induction Machines - Comparison // Procedia Engineering. 2017. Vol. 192. P. 851-856. DOI: 10.1016/j.proeng.2017.06.147.",
        "Gadoue S. M., Giaouris D., Finch J. W. Artificial intelligence-based speed control of DTC induction motor drives - A comparative study // Electric Power Systems Research. 2009. Vol. 79, no. 1. P. 210-219. DOI: 10.1016/j.epsr.2008.05.024.",
        "Szoke E., Szabo C., Pintilie L.-N. Artificial Intelligence-Based Sensorless Control of Induction Motors with Dual-Field Orientation // Applied Sciences. 2025. Vol. 15, no. 16. Art. 8919. DOI: 10.3390/app15168919.",
        "Mouodo L. V. A., Axaopoulos P., Patrice N. N. T., Abdelkerim A. A., Kibong M. T., Mouzong M. P., Tamba J. G. Design of an optimal vector control of an induction motor for electric vehicles // Results in Engineering. 2026. Vol. 30. Art. 110164. DOI: 10.1016/j.rineng.2026.110164.",
        "STMicroelectronics. STM32G431x6/x8/xB Data Sheet [Электронный ресурс]. URL: https://www.st.com/resource/en/datasheet/stm32g431rb.pdf (дата обращения: 25.08.2026).",
        "STMicroelectronics. STEVAL-IPM15B product page [Электронный ресурс]. URL: https://www.st.com/en/evaluation-tools/steval-ipm15b.html (дата обращения: 25.08.2026).",
        "STMicroelectronics. UM2014. 1500 W motor control power board based on STGIB15CH60TS-L SLLIMM 2nd series IPM [Электронный ресурс]. URL: https://www.st.com/resource/en/user_manual/um2014-1500-w-motor-control-power-board-based-on-stgib15ch60tsl-sllimm-2nd-series-ipm-stmicroelectronics.pdf (дата обращения: 25.08.2026).",
        "STMicroelectronics. UM3030. Getting started with the X-NUCLEO-IHM09M2 motor-control connector expansion board [Электронный ресурс]. URL: https://www.st.com/resource/en/user_manual/um3030-getting-started-with-the-xnucleoihm09m2-motor-control-connector-expansion-board-for-stm32-nucleo-stmicroelectronics.pdf (дата обращения: 25.08.2026).",
        "STMicroelectronics. X-CUBE-MCSDK Motor Control Software Development Kit [Электронный ресурс]. URL: https://www.st.com/en/embedded-software/x-cube-mcsdk.html (дата обращения: 25.08.2026).",
        "STMicroelectronics. STM32 Motor Control AC induction motor wiki: V/F and LSO-FOC [Электронный ресурс]. URL: https://wiki.st.com/stm32mcu/wiki/STM32_MC_ACIM (дата обращения: 25.08.2026).",
        "Arduino. UNO Q hardware documentation [Электронный ресурс]. URL: https://docs.arduino.cc/hardware/uno-q (дата обращения: 25.08.2026).",
        "IEK. Электродвигатели асинхронные трёхфазные серии АИР. Каталог [Электронный ресурс]. URL: https://cdn-01.iek.ru/media/original/78c31060549010eaa9a1bb1f4f6d2d8cb3155062b138bfd13478f55521c38a47.pdf (дата обращения: 25.08.2026).",
        "IEC 61800-5-2:2016. Adjustable speed electrical power drive systems. Safety requirements. Functional.",
        "IEC 61800-3:2017. Adjustable speed electrical power drive systems. EMC requirements and specific test methods.",
        "IEEE Std 112-2017. Standard Test Procedure for Polyphase Induction Motors and Generators.",
        "ISO 26262-6:2018. Road vehicles. Functional safety. Product development at the software level. Используется методически для принципов safety case; не является отраслевым стандартом железнодорожного привода.",
    ]
    for idx, ref in enumerate(refs, start=1):
        p = doc.add_paragraph(style="Bibliography")
        r = p.add_run(f"{idx}. {ref}")
        set_run_font(r, size=12)


def add_appendices(doc: Document, figs: dict[str, Path]) -> None:
    doc.add_heading("ПРИЛОЖЕНИЕ А\nКаталожный и расчётный профиль двигателя АИР56В2", level=1)
    profile_path = ROOT / "docs" / "mcsdk_acim_motor_profile.iek_air56b2_catalog_operator_confirmed_vf_candidate.json"
    profile = json.loads(profile_path.read_text(encoding="utf-8"))
    profile_fields = [
        "schema",
        "motor_label",
        "pole_pairs",
        "rated_line_voltage_v",
        "controller_equivalent_phase_voltage_v",
        "rated_current_a",
        "rated_frequency_hz",
        "rated_speed_rpm",
        "rated_power_w",
        "connection",
    ]
    rows = [[key, str(profile[key])] for key in profile_fields]
    add_table(doc, "Таблица А.1 - Содержимое зафиксированного расчётного профиля", ["Поле", "Значение"], rows, [3500, 6138], font_size=9)
    add_body(doc, "Примечание — Rs, Rr, Lls, Llr, Lm и J не измерены для целевого АИР56В2 и поэтому не включены в таблицу. В импортированном IOC сохранились демонстрационные значения Rs=2,85; Rr=0,7; Lls=Llr=0,003; Lm=0,1485, тогда как сгенерированный acim_motor_parameters.h содержит RS=0,35. Это выявленное противоречие подтверждает, что ни одно из значений нельзя считать параметром целевого двигателя; настройка FOC запрещена до идентификации.")
    add_noindent(doc, "Каталог проекта: mcsdk_reference/.../AIR56B2_025KW_220V_DELTA_NAMEPLATE_VF_NOT_FOR_HV. Слово NAMEPLATE в этом унаследованном идентификаторе каталога не является доказательством происхождения профиля. Маркер NOT_FOR_HV не удаляется до выполнения gate-критериев приложения Г.", align=WD_ALIGN_PARAGRAPH.LEFT, size=12)

    doc.add_heading("ПРИЛОЖЕНИЕ Б\nТаблица соединения UART верхнего уровня", level=1)
    add_table(doc, "Таблица Б.1 - Соединения через ISO7721", ["Сигнал", "UNO Q", "Изолятор", "NUCLEO-G431RB"], [
        ["Команда в привод", "D1 / Serial1 TX", "A in → A out", "PB7 / USART1_RX"],
        ["Ответ", "D0 / Serial1 RX", "B out ← B in", "PB6 / USART1_TX"],
        ["Питание стороны UNO", "3,3 V / GND UNO", "VCCA / GNDA", "не соединять"],
        ["Питание стороны Nucleo", "не соединять", "VCCB / GNDB", "3,3 V / GND Nucleo"],
        ["Программный ESTOP", "flags=0x02", "канал UART", "MC_StopMotor1(); latch fault"],
        ["Аппаратный E-stop", "отдельная НЗ цепь", "не через UART", "ещё не реализован"],
    ], [2200, 2400, 2200, 2838], font_size=9)
    add_figure(doc, figs["uart"], "Рисунок Б.1 - Формат командного кадра интерфейса")

    doc.add_heading("ПРИЛОЖЕНИЕ В\nФорма протокола стендового испытания", level=1)
    add_table(doc, "Таблица В.1 - Идентификация опыта", ["Поле", "Значение"], [
        ["Дата и оператор", "[заполнить]"],
        ["Сценарий / номер повтора", "[заполнить]"],
        ["Алгоритм и версия", "[заполнить]"],
        ["SHA-256 Nucleo", "[заполнить]"],
        ["SHA-256 UNO Q", "[заполнить]"],
        ["Профиль двигателя", "[заполнить]"],
        ["Udc / ограничение тока", "[заполнить]"],
        ["Температура до / после", "[заполнить]"],
        ["Калибровка датчиков", "[ссылка на протокол]"],
        ["Результат gate", "PASS / FAIL"],
        ["Замечания", "[заполнить]"],
    ], [3100, 6538], font_size=10)
    add_table(doc, "Таблица В.2 - Фиксация результатов", ["Метрика", "Единица", "Значение", "Неопределённость"], [
        ["КПД", "%", "", ""],
        ["IAE скорости", "об", "", ""],
        ["RMS тока", "А", "", ""],
        ["Максимальный ток", "А", "", ""],
        ["Пульсации момента", "%", "", ""],
        ["Коммутационные события", "1/с", "", ""],
        ["Fallback", "число", "", ""],
        ["Fault", "код", "", ""],
    ], [3000, 1700, 2300, 2638], font_size=10)

    doc.add_heading("ПРИЛОЖЕНИЕ Г\nЧек-лист допуска к силовому запуску", level=1)
    checks = [
        "Фото шильдика и схема перемычек Δ подтверждены.",
        "Измерены сопротивления обмоток; нет замыкания на корпус.",
        "Параметры профиля соответствуют конкретному двигателю.",
        "Предохранители, корпус, PE и разряд звена проверены.",
        "Предзаряд реализован и HIL-проверен.",
        "PWM запрещён до готовности Udc.",
        "Break/OCP/E-stop останавливают привод без UART.",
        "Timeout UART 300 мс проверен.",
        "Фазировка токов и фаз двигателя проверена на низком напряжении.",
        "Источник ограничен по напряжению и току.",
        "Осциллограф и пробники рассчитаны на категорию измерения.",
        "Назначен наблюдающий; зона ограждена; аварийное отключение доступно.",
        "Разрешение на текущий этап подписано ответственным лицом.",
    ]
    add_table(doc, "Таблица Г.1 - Контроль перед подачей высокого напряжения", ["№", "Проверка", "Отметка / подпись"], [[str(i), item, ""] for i, item in enumerate(checks, 1)], [700, 7038, 1900], font_size=9)

    doc.add_heading("ПРИЛОЖЕНИЕ Д\nМатрица трассируемости задач и доказательств", level=1)
    add_table(doc, "Таблица Д.1 - Связь задач, программных объектов и подтверждений", ["Задача", "Объект проекта", "Текущее доказательство", "Будущее доказательство"], [
        ["V/f, 0 < f ≤ 50 Гц", "MCSDK AIR56B2 project", "Build PASS + числовая согласованность", "Стенд S3"],
        ["Безопасный UART", "main.c + UNOQ_MOTOR.ino", "Статический contract PASS", "Runtime timeout + corrupted frames"],
        ["Аппаратный останов", "break/OCP + будущий E-stop", "Break предусмотрен; E-stop отсутствует", "Fault injection/HIL"],
        ["FOC", "Методика главы 2", "Математическая модель", "S4 после ID"],
        ["ИИ-адаптация", "historical_host_release", "Host scenarios", "Shadow + A/B"],
        ["ИИ-ШИМ", "SNH-PWM prototype", "Gateway fault injection", "MCU timing + HIL"],
        [">50 Гц", "Field-weakening plan", "Расчётная область", "Механический допуск"],
    ], [2000, 2500, 2600, 2538], font_size=9)
    add_body(doc, "Матрица трассируемости используется как индекс доказательств: каждой задаче ставятся в соответствие исходный объект, уже полученный результат и эксперимент, который ещё требуется выполнить. Строка считается закрытой только при наличии исходных данных, версии алгоритма, протокола измерения и воспроизводимого отчёта.")
    add_body(doc, "Слова PASS в колонке текущего доказательства относятся только к программной проверке указанного уровня. Они не означают допуск к подаче высокого напряжения и не заменяют будущие доказательства из последней колонки.")
    add_body(doc, "При изменении двигателя, платы, версии MCSDK, распиновки, протокола или параметров защиты формируется новая запись матрицы и новый комплект хэшей. Результаты предыдущей конфигурации не переносятся автоматически.")


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    figs = make_figures()
    doc = Document()
    configure_document(doc)
    props = doc.core_properties
    props.title = "Рабочая редакция диссертации по специальности 2.9.3"
    props.subject = "Асинхронный тяговый электропривод, FOC, ограниченная интеллектуальная адаптация"
    props.author = "Соискатель не указан; рабочая редакция проекта MIC_AI"
    props.keywords = "2.9.3; асинхронный привод; FOC; AI; SVPWM; STM32G431; ПГУПС"
    props.comments = "Рабочая редакция. Стендовые результаты не сфабрикованы и должны быть заполнены после испытаний."
    now = datetime.now(timezone.utc)
    props.created = now
    props.modified = now

    add_title_page(doc)
    add_status_page(doc)
    add_toc_and_abbreviations(doc)
    add_intro(doc, figs)
    add_chapter1(doc)
    add_chapter2(doc, figs)
    add_chapter3(doc, figs)
    add_chapter4(doc, figs)
    add_chapter5(doc, figs)
    add_conclusion(doc)
    add_illustrative_material(doc)
    add_references(doc)
    add_appendices(doc, figs)

    normalize_text_appearance(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
