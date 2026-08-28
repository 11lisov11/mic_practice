#!/usr/bin/env python3
from __future__ import annotations

import argparse
from collections import defaultdict
from datetime import date
import json
import math
from pathlib import Path
import statistics
from typing import Any, Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS = ROOT / "artifacts"
OUTPUT = ROOT / "output" / "documents"
ASSETS = OUTPUT / "air56b2_report_assets"
REPORT_PATH = OUTPUT / "AIR56B2_научно_технический_отчет_2026-08-28.docx"
PROTOCOL_PATH = OUTPUT / "AIR56B2_программа_аппаратной_верификации_2026-08-28.docx"
FONT_REGULAR = Path(r"C:\Windows\Fonts\times.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\timesbd.ttf")
BLACK = RGBColor(0, 0, 0)
CONTENT_WIDTH_DXA = 9354
TABLE_INDENT_DXA = 120


def _read(name: str) -> dict[str, Any]:
    path = ARTIFACTS / name
    if not path.is_file():
        raise FileNotFoundError(path)
    return json.loads(path.read_text(encoding="utf-8"))


def _font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def _centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font: ImageFont.FreeTypeFont) -> None:
    left, top, right, bottom = box
    bounds = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(
        ((left + right - width) / 2, (top + bottom - height) / 2),
        text,
        font=font,
        fill="black",
        spacing=6,
        align="center",
    )


def build_pipeline_figure(path: Path) -> None:
    image = Image.new("L", (1800, 980), "white")
    draw = ImageDraw.Draw(image)
    title = _font(FONT_BOLD, 44)
    body = _font(FONT_REGULAR, 31)
    small = _font(FONT_REGULAR, 25)
    draw.text((900, 38), "Воспроизводимый контур исследования AIR56B2", font=title, fill="black", anchor="ma")
    boxes = [
        (80, 180, 380, 390, "Паспорт\nAIR56B2\nF1 / F1S"),
        (470, 180, 770, 390, "Ансамбль\nF2 / F3\n256 образцов"),
        (860, 105, 1210, 315, "Регуляторы\nV/f, FOC,\nsensorless"),
        (860, 400, 1210, 610, "Оптимизация\nпотерь и\nтепловой prior"),
        (1300, 250, 1710, 470, "Disjoint benchmark\ntrain / validation /\nholdout"),
        (1300, 610, 1710, 820, "Manifest v6\nхеши, ворота,\nrelease=false"),
    ]
    for left, top, right, bottom, text in boxes:
        draw.rounded_rectangle((left, top, right, bottom), radius=8, outline="black", width=4, fill=245)
        _centered_text(draw, (left + 15, top + 10, right - 15, bottom - 10), text, body)
    arrows = [
        ((380, 285), (470, 285)),
        ((770, 285), (860, 210)),
        ((770, 285), (860, 505)),
        ((1210, 210), (1300, 330)),
        ((1210, 505), (1300, 390)),
        ((1505, 470), (1505, 610)),
    ]
    for start, end in arrows:
        draw.line((start, end), fill="black", width=5)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        for offset in (2.65, -2.65):
            point = (
                end[0] + 22 * math.cos(angle + offset),
                end[1] + 22 * math.sin(angle + offset),
            )
            draw.line((end, point), fill="black", width=5)
    draw.text(
        (900, 910),
        "Все численные результаты относятся к моделированию; аппаратная идентификация остаётся обязательной.",
        font=small,
        fill="black",
        anchor="ma",
    )
    image.save(path, dpi=(220, 220))


def build_loss_heatmap(path: Path, study: dict[str, Any]) -> None:
    grouped: dict[tuple[float, float], list[float]] = defaultdict(list)
    for row in study["cases"]:
        grouped[(float(row["speed_fraction"]), float(row["torque_fraction"]))].append(
            float(row["loss_saving_pct"])
        )
    speeds = sorted({key[0] for key in grouped})
    torques = sorted({key[1] for key in grouped})
    values = {(speed, torque): statistics.median(grouped[(speed, torque)]) for speed in speeds for torque in torques}
    minimum = min(values.values())
    maximum = max(values.values())
    image = Image.new("L", (1700, 1180), "white")
    draw = ImageDraw.Draw(image)
    title = _font(FONT_BOLD, 42)
    axis = _font(FONT_REGULAR, 30)
    cell_font = _font(FONT_BOLD, 31)
    draw.text((850, 40), "Медианное снижение расчётных потерь, %", font=title, fill="black", anchor="ma")
    left, top, cell_w, cell_h = 300, 165, 240, 170
    for row_index, torque in enumerate(reversed(torques)):
        y = top + row_index * cell_h
        draw.text((left - 30, y + cell_h / 2), f"{torque:.2f}", font=axis, fill="black", anchor="rm")
        for col_index, speed in enumerate(speeds):
            x = left + col_index * cell_w
            value = values[(speed, torque)]
            fraction = 0.0 if maximum == minimum else (value - minimum) / (maximum - minimum)
            shade = int(242 - 115 * fraction)
            draw.rectangle((x, y, x + cell_w, y + cell_h), fill=shade, outline="black", width=3)
            text_fill = "white" if shade < 160 else "black"
            draw.text((x + cell_w / 2, y + cell_h / 2), f"{value:.1f}", font=cell_font, fill=text_fill, anchor="mm")
    bottom = top + len(torques) * cell_h
    for col_index, speed in enumerate(speeds):
        x = left + col_index * cell_w + cell_w / 2
        draw.text((x, bottom + 24), f"{speed:.2f}", font=axis, fill="black", anchor="ma")
    draw.text((left + len(speeds) * cell_w / 2, bottom + 95), "Скорость, о.е.", font=axis, fill="black", anchor="ma")
    draw.text((90, top + len(torques) * cell_h / 2), "Момент, о.е.", font=axis, fill="black", anchor="mm")
    image.save(path, dpi=(220, 220))


def build_policy_figure(path: Path, benchmark: dict[str, Any]) -> None:
    rows = benchmark["holdout_rows"]
    points = [
        (float(row["loss_saving_vs_fixed_pct"]), float(row["optimality_gap_pct"]))
        for row in rows
    ]
    x_min = min(0.0, min(value[0] for value in points)) * 1.10
    x_max = max(value[0] for value in points) * 1.05
    y_min = min(0.0, min(value[1] for value in points)) * 1.10
    y_max = max(value[1] for value in points) * 1.12
    image = Image.new("L", (1700, 1050), "white")
    draw = ImageDraw.Draw(image)
    title = _font(FONT_BOLD, 42)
    axis = _font(FONT_REGULAR, 29)
    draw.text((850, 38), "Holdout: выигрыш политики и разрыв от оптимума", font=title, fill="black", anchor="ma")
    left, right, top, bottom = 270, 1570, 145, 820
    draw.line((left, bottom, right, bottom), fill="black", width=4)
    draw.line((left, top, left, bottom), fill="black", width=4)
    for tick in range(0, 6):
        x_value = x_min + (x_max - x_min) * tick / 5
        x = left + (right - left) * tick / 5
        draw.line((x, bottom, x, bottom + 12), fill="black", width=3)
        draw.text((x, bottom + 20), f"{x_value:.0f}", font=axis, fill="black", anchor="ma")
        y_value = y_min + (y_max - y_min) * tick / 5
        y = bottom - (bottom - top) * tick / 5
        draw.line((left - 12, y, left, y), fill="black", width=3)
        y_label_value = 0.0 if abs(y_value) < 0.05 else y_value
        draw.text((left - 22, y), f"{y_label_value:.1f}", font=axis, fill="black", anchor="rm")
    for saving, gap in points:
        x = left + (right - left) * (saving - x_min) / max(x_max - x_min, 1e-9)
        y = bottom - (bottom - top) * (gap - y_min) / max(y_max - y_min, 1e-9)
        draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill=80)
    zero_x = left + (right - left) * (0.0 - x_min) / max(x_max - x_min, 1e-9)
    draw.line((zero_x, top, zero_x, bottom), fill=150, width=2)
    draw.text(((left + right) / 2, bottom + 95), "Снижение потерь относительно фиксированного потока, %", font=axis, fill="black", anchor="ma")
    y_label = Image.new("L", (520, 60), "white")
    y_draw = ImageDraw.Draw(y_label)
    y_draw.text((260, 30), "Разрыв от оптимума, %", font=axis, fill="black", anchor="mm")
    rotated = y_label.rotate(90, expand=True, fillcolor="white")
    image.paste(rotated, (28, int((top + bottom - rotated.height) / 2)))
    image.save(path, dpi=(220, 220))


def build_common_benchmark_figure(path: Path, benchmark: dict[str, Any]) -> None:
    method_order = (
        ("fixed", "Фиксированный поток"),
        ("classical_optimum", "Классический оптимум"),
        ("neural_policy", "Нейронная политика"),
        ("extremum_search", "Поиск экстремума"),
        ("guarded_lut", "LUT с fallback"),
    )
    image = Image.new("L", (1700, 1080), "white")
    draw = ImageDraw.Draw(image)
    title = _font(FONT_BOLD, 42)
    axis = _font(FONT_REGULAR, 29)
    label = _font(FONT_REGULAR, 31)
    note = _font(FONT_REGULAR, 25)
    draw.text((850, 38), "Парное сравнение на едином holdout (600 режимов)", font=title, fill="black", anchor="ma")
    left, right, top, bottom = 490, 1570, 155, 850
    x_max = 18.0
    draw.line((left, bottom, right, bottom), fill="black", width=4)
    for tick in range(0, 7):
        value = x_max * tick / 6
        x = left + (right - left) * value / x_max
        draw.line((x, bottom, x, bottom + 12), fill="black", width=3)
        draw.text((x, bottom + 20), f"{value:.0f}", font=axis, fill="black", anchor="ma")
    row_height = (bottom - top) / len(method_order)
    for index, (key, name) in enumerate(method_order):
        y = top + row_height * (index + 0.5)
        metrics = benchmark["methods"][key]["saving_vs_fixed_pct"]
        mean = float(metrics["mean"])
        median = float(metrics["median"])
        ci_low, ci_high = (float(value) for value in metrics["mean_cluster_bootstrap_95_ci"])
        x_low = left + (right - left) * max(0.0, ci_low) / x_max
        x_high = left + (right - left) * min(x_max, ci_high) / x_max
        x_mean = left + (right - left) * max(0.0, mean) / x_max
        x_median = left + (right - left) * max(0.0, median) / x_max
        draw.text((left - 28, y), name, font=label, fill="black", anchor="rm")
        if key == "fixed":
            draw.ellipse((x_mean - 7, y - 7, x_mean + 7, y + 7), fill="black")
            continue
        draw.line((x_low, y, x_high, y), fill="black", width=5)
        draw.line((x_low, y - 12, x_low, y + 12), fill="black", width=3)
        draw.line((x_high, y - 12, x_high, y + 12), fill="black", width=3)
        draw.ellipse((x_mean - 10, y - 10, x_mean + 10, y + 10), fill="white", outline="black", width=4)
        draw.line((x_median - 8, y - 10, x_median + 8, y + 10), fill="black", width=3)
        draw.line((x_median - 8, y + 10, x_median + 8, y - 10), fill="black", width=3)
    draw.text(((left + right) / 2, bottom + 95), "Снижение потерь относительно фиксированного потока, %", font=axis, fill="black", anchor="ma")
    draw.text((850, 1010), "Круг — среднее; горизонтальный интервал — кластерный 95%-й ДИ; крест — медиана.", font=note, fill="black", anchor="ma")
    image.save(path, dpi=(220, 220))


def _set_run_font(run, *, size: float = 14, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK


def _configure_document(document: Document, running_title: str) -> None:
    section = document.sections[0]
    section.different_first_page_header_footer = True
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.5
    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 12, 6),
        ("Heading 3", 14, 10, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(running_title)
    _set_run_font(run, size=9)
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    section.first_page_header.paragraphs[0].clear()
    section.first_page_footer.paragraphs[0].clear()


def _title_page(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    for line in (
        "Федеральное государственное бюджетное образовательное учреждение",
        "высшего образования",
        "«Петербургский государственный университет путей сообщения",
        "Императора Александра I»",
    ):
        run = paragraph.add_run(line + "\n")
        _set_run_font(run, size=12)
    document.add_paragraph()
    document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(20)
    run = paragraph.add_run(title)
    _set_run_font(run, size=16, bold=True)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    _set_run_font(run, size=14)
    for _ in range(7):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in ("Проект: MIC AI", "Объект: АИР56В2, 0,25 кВт, 220 В, Δ", "Версия результатов: 28.08.2026"):
        run = paragraph.add_run(line + "\n")
        _set_run_font(run, size=14)
    for _ in range(4):
        document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Санкт-Петербург\n2026")
    _set_run_font(run, size=14)
    document.add_page_break()


def _add_body(document: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        _set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        _set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        _set_run_font(run)


def _set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)


def _shade(cell, fill: str = "D9D9D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _add_table(document: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _shade(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        _set_run_font(run, size=11, bold=True)
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_tr_pr.append(repeat_header)
    for values in rows:
        row = table.add_row()
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        row_pr.append(cant_split)
        cells = row.cells
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            _set_run_font(run, size=11)
    _set_table_geometry(table, widths_dxa)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)


def _add_figure(document: Document, path: Path, caption: str, number: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Mm(160))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.keep_with_next = True
    caption_run = caption_paragraph.add_run(f"Рисунок {number} — {caption}")
    _set_run_font(caption_run, size=12)


def build_report(output_path: Path, assets: dict[str, Path], data: dict[str, dict[str, Any]]) -> None:
    manifest = data["manifest"]
    loss = data["loss"]
    sensorless = data["sensorless"]
    policy = data["policy"]
    common = data["common"]
    vf = data["vf"]
    matrix = data["matrix"]
    faults = data["faults"]
    encoder = data["encoder"]
    document = Document()
    _configure_document(document, "MIC AI | AIR56B2 | научно-технический отчёт")
    _title_page(
        document,
        "НАУЧНО-ТЕХНИЧЕСКИЙ ОТЧЁТ",
        "по программно-модельной подготовке электропривода AIR56B2",
    )
    document.add_heading("Аннотация", level=1)
    _add_body(
        document,
        "В отчёте зафиксирован завершённый объём программно-модельной подготовки экспериментального электропривода с асинхронным двигателем AIR56B2. Реализованы паспортно-ограниченный ансамбль моделей, скалярное и векторное управление, независимая бездатчиковая оценка скорости, модель потерь и температур, классическая оптимизация тока намагничивания, нейронная аппроксимация оптимума и единый воспроизводимый конвейер. Все приведённые результаты относятся к моделированию. Они не заменяют измерение параметров двигателя, низковольтную наладку и аппаратную проверку защит.",
    )
    _add_table(
        document,
        ["Статус", "Значение"],
        [
            ["Программно-модельный пакет", manifest["status"]],
            ["Готовность к аппаратной идентификации", "да" if manifest["pc_research_ready_for_hardware_identification"] else "нет"],
            ["Разрешение силового аппаратного запуска", "нет"],
            ["Флаг hardware_release_ready", str(manifest["hardware_release_ready"]).lower()],
            ["Автоматические проверки", "220 тестов: 169 + 34 + 17"],
        ],
        [4200, 5154],
    )
    document.add_heading("1 Цель и границы исследования", level=1)
    _add_body(
        document,
        "Цель этапа — подготовить на персональном компьютере проверяемую основу для последовательного перехода от открытого закона U/f к FOC, затем к адаптивной оптимизации и бездатчиковому управлению. Критерием завершения программной части принято наличие воспроизводимых артефактов, раздельных train/validation/holdout выборок, формализованных ворот безопасности и явного перечня проверок, которые невозможно выполнить без реальных плат и двигателя.",
    )
    _add_figure(document, assets["pipeline"], "Структура воспроизводимого программно-модельного контура", 1)
    document.add_heading("2 Исходные данные и иерархия моделей", level=1)
    _add_table(
        document,
        ["Параметр", "Паспортное значение"],
        [
            ["Тип двигателя", "AIR56B2"],
            ["Номинальная мощность", "0,25 кВт"],
            ["Напряжение и соединение", "220 В, Δ"],
            ["Номинальный ток", "1,24 А"],
            ["Частота", "50 Гц"],
            ["Номинальная скорость", "2720 об/мин"],
            ["cos φ / КПД", "0,78 / 0,68"],
            ["Число пар полюсов", "1"],
        ],
        [4300, 5054],
    )
    _add_body(
        document,
        "Уровень F1 представляет одно-клеточную схему замещения, ограниченную паспортной рабочей точкой. Уровень F1S вводит феноменологическую поправку высокоскользящих потерь по паспортным кратностям пускового и максимального момента. Уровень F2 добавляет температурные коэффициенты, насыщение и механические потери как ограниченные приоры. Уровень F3 добавляет неидеальности инвертора, АЦП и AS5600. Электрические сопротивления, индуктивности, инерция и тепловые постоянные не объявляются паспортными или измеренными.",
    )
    document.add_heading("3 Реализованные методы", level=1)
    _add_body(document, "Скалярный baseline. Реализован закон U/f с ограничением напряжения, контролем токовой перегрузки и защитным шлюзом. Проверка охватывает четыре частоты: 5, 15, 30 и 50 Гц.", bold_lead="Скалярный baseline.")
    _add_body(document, "Векторное управление. Oracle-FOC используется только как верхняя модельная граница и явно принимает полный simulated state. Отдельный encoder-observer FOC использует задержанные квантованные токи, напряжение звена и AS5600, но не истинные поток, угол или скорость.", bold_lead="Векторное управление.")
    _add_body(document, "Бездатчиковый наблюдатель. Новый voltage/flux/slip observer принимает только восстановленное αβ-напряжение и измеренные токи. Валидация выполнена на независимой RK4-модели с состояниями тока статора и потока ротора.", bold_lead="Бездатчиковый наблюдатель.")
    _add_body(document, "Оптимизация и ИИ. Классический constrained baseline минимизирует сумму медных, магнитных, механических и инверторных потерь по i_d. Нейронная политика является supervised-distillation этого optimum, а не аппаратно обученным регулятором.", bold_lead="Оптимизация и ИИ.")
    document.add_heading("4 Численные результаты", level=1)
    _add_table(
        document,
        ["Проверка", "Результат"],
        [
            ["V/f, базовая выборка", f"{vf['summary']['passed_trials']}/{vf['sample_count']} PASS"],
            ["V/f, рабочая матрица", f"{matrix['summary']['passed_trials']}/{matrix['total_trial_count']} PASS"],
            ["Защитный шлюз", f"{faults['summary']['passed_fault_cases']}/{faults['total_fault_case_count']} PASS"],
            ["Encoder-observer FOC", f"{encoder['summary']['passed_trials']}/{encoder['sample_count']} PASS"],
            ["Loss baseline", f"{loss['case_count_comparable']}/{loss['case_count_expected']} допустимых точек"],
            ["Sensorless independent plant", f"средняя ошибка {sensorless['validation_summary']['mean_abs_speed_error_rad_s']:.2f} рад/с"],
            ["Neural policy holdout", f"{policy['holdout_summary']['case_count']} точек, 0 нарушений"],
        ],
        [5000, 4354],
    )
    _add_body(
        document,
        f"Для классического baseline медианное расчётное снижение суммарных потерь относительно фиксированного номинального потока составило {loss['summary']['loss_saving_pct_median']:.2f} %, среднее — {loss['summary']['loss_saving_pct_mean']:.2f} %. На всех сопоставимых точках оптимизатор не ухудшил исходную уставку.",
    )
    _add_figure(document, assets["loss"], "Медианный выигрыш классической оптимизации по ансамблю F2/F3", 2)
    _add_body(
        document,
        f"На независимом sensorless holdout для 15, 30 и 45 Гц средняя абсолютная ошибка скорости составила {sensorless['validation_summary']['mean_abs_speed_error_rad_s']:.2f} рад/с, худшая средняя ошибка — {sensorless['validation_summary']['worst_mean_abs_speed_error_rad_s']:.2f} рад/с, а средняя относительная ошибка — {sensorless['validation_summary']['mean_relative_speed_error_pct']:.2f} %. Диагностика 5 Гц сохраняется как неподтверждённая для аппаратного перехода: voltage-model наблюдатель принципиально чувствителен к смещению и ошибке R_s на малой частоте.",
    )
    _add_body(
        document,
        f"Нейронная политика проверена на {policy['holdout_summary']['case_count']} ранее не использованных точках. Медианный выигрыш относительно фиксированного потока равен {policy['holdout_summary']['loss_saving_vs_fixed_pct_median']:.2f} %, медианный разрыв от constrained optimum — {policy['holdout_summary']['optimality_gap_pct_median']:.2f} %, максимальный — {policy['holdout_summary']['optimality_gap_pct_max']:.2f} %. Нарушений ограничений по току и напряжению не зарегистрировано.",
    )
    _add_figure(document, assets["policy"], "Распределение качества нейронной политики на holdout", 3)
    _add_body(
        document,
        f"На едином парном holdout из {common['case_count']} режимов классический constrained optimum дал среднее снижение расчётных потерь {common['methods']['classical_optimum']['saving_vs_fixed_pct']['mean']:.2f} %, нейронная политика — {common['methods']['neural_policy']['saving_vs_fixed_pct']['mean']:.2f} %, ограниченный поиск экстремума — {common['methods']['extremum_search']['saving_vs_fixed_pct']['mean']:.2f} %, а LUT с fallback — {common['methods']['guarded_lut']['saving_vs_fixed_pct']['mean']:.2f} %. Нейронная политика ухудшила фиксированный baseline в {common['methods']['neural_policy']['worse_than_fixed_count']} из {common['case_count']} точек; худшее изменение составило {common['methods']['neural_policy']['saving_vs_fixed_pct']['minimum']:.2f} %. Поэтому активное применение actor без надзорного контура и проверяемого fallback не допускается.",
    )
    _add_figure(document, assets["common"], "Парное сравнение стратегий управления на общем holdout", 4)
    document.add_heading("5 Воспроизводимость", level=1)
    _add_body(
        document,
        "Полный пакет пересобирается командой powershell -ExecutionPolicy Bypass -File tools\\run_air56b2_research_pipeline.ps1 -Full. Конвейер фиксирует Python, PyTorch, CUDA, драйвер и RTX 5070; пересоздаёт ансамбль 256 образцов; прогоняет тесты, train/validation/holdout, sensorless, policy и общий парный benchmark; затем формирует manifest v6 с SHA-256 каждого канонического JSON. Два независимых CUDA-smoke прогона FOC+PPO с seed 560225 дали побитно одинаковые actor, critic и metrics.",
    )
    _add_table(
        document,
        ["Артефакт", "Назначение"],
        [
            ["air56b2_research_manifest.json", "единая цепочка доказательств и ворота"],
            ["air56b2_fidelity_bundle.json", "F1/F1S/F2/F3, 256 образцов"],
            ["air56b2_loss_optimization_\nstudy.json", "480 рабочих точек loss baseline"],
            ["air56b2_sensorless_\nindependent_plant_study.json", "sensorless train/validation"],
            ["air56b2_policy_benchmark.json", "disjoint holdout нейронной политики"],
            ["air56b2_common_control_\nbenchmark.json", "парное сравнение стратегий и 95%-е ДИ"],
            ["air56b2_id_policy_actor.pt", "simulation-only actor checkpoint"],
            ["air56b2_id_ref_lut.h", "LUT, аппаратный release отключён"],
        ],
        [4300, 5054],
    )
    document.add_heading("6 Научная интерпретация", level=1)
    _add_body(
        document,
        "Полученные результаты поддерживают основную исследовательскую гипотезу: рабочий регулятор может не вычислять в каждом цикле полную параметрическую модель конкретного двигателя. Модель и constrained optimizer используются на этапе генерации и проверки политики, а в реальном контуре компактный actor или LUT формирует корректирующую уставку по измеряемой динамике и доступным оценкам. Это не абсолютный model-free подход: физическая модель сохраняется как обучающий, ограничивающий и верификационный слой.",
    )
    _add_body(
        document,
        "Наиболее сильный научный результат текущего этапа — не величина модельного процента экономии сама по себе, а разделение ролей: oracle-модель задаёт верхнюю границу; независимый объект выявляет переобучение к реализации; classical optimum создаёт интерпретируемый baseline; holdout показывает аппроксимационную способность; hardware gate не позволяет перенести модельное утверждение на двигатель без измерений.",
    )
    document.add_heading("7 Ограничения и незакрытые риски", level=1)
    limitations = [
        "Параметры R_s, R_r, L_σ, L_m, J и B пока не измерены на экземпляре AIR56B2.",
        "F1S является феноменологической high-slip поправкой, а не идентифицированной двухклеточной схемой.",
        "Потери в стали, switching-time equivalent и тепловые постоянные остаются simulation priors.",
        "Низкоскоростной sensorless режим, регенерация и ослабление поля не разрешены для аппаратного release.",
        "Actor и LUT не должны напрямую управлять ШИМ: обязательны saturation, current/voltage/temperature supervisor и fallback.",
        "Силовая часть 310 В постоянного тока требует физического разрыва с ПК, аппаратного аварийного останова и безопасного разряда звена.",
    ]
    for index, item in enumerate(limitations, 1):
        _add_body(document, f"{index}. {item}")
    document.add_heading("8 Выводы", level=1)
    _add_body(
        document,
        "Программно-модельный этап завершён в объёме, достаточном для перехода к аппаратной идентификации и поэтапной стендовой проверке. Подготовлены скалярный baseline, encoder-observer FOC, экспериментальный sensorless observer, loss baseline, нейронная policy и LUT, а также единый manifest. Непосредственная высоковольтная эксплуатация не разрешена: следующий доказательный уровень начинается только после появления плат, измерения параметров и прохождения программы аппаратной верификации.",
    )
    document.add_heading("Приложение А. Команда полного прогона", level=1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run("powershell -ExecutionPolicy Bypass -File tools\\run_air56b2_research_pipeline.ps1 -Full")
    _set_run_font(run, size=11)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_protocol(output_path: Path, data: dict[str, dict[str, Any]]) -> None:
    document = Document()
    _configure_document(document, "MIC AI | AIR56B2 | программа аппаратной верификации")
    _title_page(
        document,
        "ПРОГРАММА АППАРАТНОЙ ВЕРИФИКАЦИИ",
        "поэтапный ввод NUCLEO-G431RB, STEVAL-IPM15B и AIR56B2",
    )
    document.add_heading("1 Назначение", level=1)
    _add_body(
        document,
        "Документ задаёт последовательность проверок после получения плат. Он запрещает переход к следующему этапу при невыполнении хотя бы одного критерия. Проверка выполняется от обесточенного монтажа к низковольтному стенду, затем к V/f, encoder FOC, sensorless и только после этого к ИИ-коррекции.",
    )
    document.add_heading("2 Обязательные условия безопасности", level=1)
    safety = [
        "Физически отключить ПК, ST-Link, USB и Saleae до подачи силового напряжения.",
        "Управление в силовом режиме выполнять автономно или по Wi-Fi через UNO Q.",
        "Использовать двухполюсное отключение сети, предохранитель, ограничение пускового тока, разряд звена и аппаратный аварийный останов.",
        "Перед касанием измерить напряжение DC-link; программный индикатор не заменяет мультиметр.",
        "Первый запуск выполнять без механической нагрузки и с ограниченным напряжением/энергией источника.",
    ]
    for index, item in enumerate(safety, 1):
        _add_body(document, f"{index}. {item}")
    document.add_heading("3 Этапы", level=1)
    stages = [
        ["0", "Монтаж без питания", "Прозвонка GND, 3,3/5/15 В, шлейфа, фаз U/V/W, отсутствия КЗ", "Все цепи соответствуют pinmap; сопротивление DC+–DC− не указывает на КЗ"],
        ["1", "Логика без IPM", "Прошивка NUCLEO и UNO Q; UART, watchdog, STOP/ESTOP, логирование", "Команды и телеметрия проходят; потеря связи переводит привод в STOP"],
        ["2", "Низковольтный DC-link", "ШИМ, dead-time, ADC offsets, порядок фаз; двигатель отключён", "Нет shoot-through; токовые каналы имеют нулевое смещение в допуске"],
        ["3", "V/f без нагрузки", "5 → 15 → 30 → 50 Гц с паузами и лимитами", "Правильное направление, ток и нагрев в пределах; останов штатный"],
        ["4", "AS5600 / encoder FOC", "Сопоставить угол, скорость, порядок фаз; включить FOC с малым моментом", "Угол непрерывен, ошибка скорости устойчива, защиты не срабатывают"],
        ["5", "Sensorless shadow", "Наблюдатель считает скорость, но не управляет; сравнение с AS5600", "Ошибка и confidence проходят пороги на 15/30/45 Гц"],
        ["6", "Sensorless takeover", "Плавный переход encoder → sensorless с fallback", "Нет скачка тока/момента; при потере confidence возврат к encoder/V/f"],
        ["7", "ИИ shadow", "Actor/LUT только предлагает i_d; classical supervisor проверяет", "0 нарушений ограничений, логи полны, отклонение от optimum объяснимо"],
        ["8", "ИИ active", "Ограниченная коррекция i_d с rate limit и аппаратными защитами", "Экономия подтверждена измерениями при неизменной скорости и моменте"],
    ]
    _add_table(document, ["№", "Режим", "Действия", "Критерий перехода"], stages, [550, 1650, 3450, 3704])
    document.add_heading("4 Минимальный набор измеряемых величин", level=1)
    _add_table(
        document,
        ["Канал", "Назначение", "Требование"],
        [
            ["I_A, I_B, I_C", "защита и FOC", "offset calibration до ШИМ; проверка знака и масштаба"],
            ["V_DC", "модуляция и защита", "сверка с мультиметром; пороги UV/OV"],
            ["AS5600", "teacher/encoder", "непрерывность угла, направление, частота обновления"],
            ["T_IPM", "защита силового модуля", "порог trip проверяется без нагрева силовой части"],
            ["ω_ref / ω_est", "качество управления", "логировать обе величины с единым timestamp"],
            ["STOP/ESTOP", "аппаратная безопасность", "должны отключать все gate-сигналы независимо от Wi-Fi"],
        ],
        [1900, 2800, 4654],
    )
    document.add_heading("5 Параметры, которые необходимо идентифицировать", level=1)
    _add_table(
        document,
        ["Параметр", "Метод", "Использование"],
        [
            ["R_s", "DC/standstill test при безопасном напряжении", "компенсация падения, observer"],
            ["R_r, τ_r", "locked-rotor или constrained transient fit", "slip и current model"],
            ["L_σ, L_m", "частотный/ступенчатый standstill test", "FOC, поток, saturation"],
            ["J, B", "run-up/coast-down", "скоростной контур и plant validation"],
            ["Потери/температура", "длительные точки 25/50/75/100 % нагрузки", "loss model и thermal supervisor"],
        ],
        [2200, 3400, 3754],
    )
    document.add_heading("6 Stop-критерии", level=1)
    stop_items = [
        "любое расхождение распиновки, направления датчика или знака тока;",
        "несоответствие V_DC измерению более установленного допуска;",
        "потеря UART/Wi-Fi без автоматического STOP;",
        "ток выше программного или аппаратного лимита;",
        "рост температуры без корректного trip;",
        "невалидный observer confidence, скачок угла или переход через нулевую скорость без fallback;",
        "любое управление actor/LUT вне supervisor envelope.",
    ]
    for index, item in enumerate(stop_items, 1):
        _add_body(document, f"{index}. {item}")
    document.add_heading("7 Протокол одного испытания", level=1)
    _add_table(
        document,
        ["Поле", "Запись"],
        [
            ["Дата / оператор", ""],
            ["Версии прошивок и manifest SHA-256", ""],
            ["Схема соединений / ревизия", ""],
            ["Напряжение DC-link", ""],
            ["Команда скорости / момент нагрузки", ""],
            ["Пиковый и установившийся ток", ""],
            ["Ошибка скорости / observer confidence", ""],
            ["Температуры", ""],
            ["Срабатывания защит", ""],
            ["Результат PASS/FAIL и причина", ""],
        ],
        [3700, 5654],
    )
    document.add_heading("8 Итоговый критерий допуска ИИ", level=1)
    _add_body(
        document,
        "ИИ допускается к активной коррекции только после аппаратной проверки V/f, encoder FOC и sensorless shadow. Эффективность сравнивают при одинаковых скорости, моменте нагрузки и временном окне. Экономия за счёт уменьшения механического выхода положительным результатом не считается.",
    )
    document.add_heading("9 Решение о допуске", level=1)
    _add_table(
        document,
        ["Контрольный пункт", "PASS/FAIL", "Примечание"],
        [
            ["Аппаратная идентификация завершена", "", ""],
            ["STOP/ESTOP и все trip-пути проверены", "", ""],
            ["V/f и encoder FOC устойчивы", "", ""],
            ["Sensorless shadow прошёл пороги", "", ""],
            ["Actor/LUT работает только через supervisor", "", ""],
            ["Логи и версии артефактов архивированы", "", ""],
        ],
        [4600, 1500, 3254],
    )
    _add_body(document, "Решение:  разрешить / запретить активную ИИ-коррекцию.")
    _add_body(document, "Ответственный: ____________________    Дата: ____________________")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build AIR56B2 final scientific report and hardware protocol")
    parser.add_argument("--report", type=Path, default=REPORT_PATH)
    parser.add_argument("--protocol", type=Path, default=PROTOCOL_PATH)
    args = parser.parse_args()
    data = {
        "manifest": _read("air56b2_research_manifest.json"),
        "loss": _read("air56b2_loss_optimization_study.json"),
        "sensorless": _read("air56b2_sensorless_independent_plant_study.json"),
        "policy": _read("air56b2_policy_benchmark.json"),
        "common": _read("air56b2_common_control_benchmark.json"),
        "vf": _read("air56b2_vf_fidelity_study.json"),
        "matrix": _read("air56b2_vf_operating_matrix.json"),
        "faults": _read("air56b2_protection_fault_matrix.json"),
        "encoder": _read("air56b2_encoder_foc_fidelity_study.json"),
    }
    ASSETS.mkdir(parents=True, exist_ok=True)
    assets = {
        "pipeline": ASSETS / "pipeline.png",
        "loss": ASSETS / "loss_heatmap.png",
        "policy": ASSETS / "policy_holdout.png",
        "common": ASSETS / "common_control_benchmark.png",
    }
    build_pipeline_figure(assets["pipeline"])
    build_loss_heatmap(assets["loss"], data["loss"])
    build_policy_figure(assets["policy"], data["policy"])
    build_common_benchmark_figure(assets["common"], data["common"])
    build_report(args.report.resolve(), assets, data)
    build_protocol(args.protocol.resolve(), data)
    print(json.dumps({"report": str(args.report.resolve()), "protocol": str(args.protocol.resolve())}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
